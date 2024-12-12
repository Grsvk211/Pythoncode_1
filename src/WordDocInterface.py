import docx
import re
import difflib
import ctypes
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor
import ExcelInterface as EI
import json
import os
from os import listdir
from os.path import isfile, join
import sys
import shutil
import QIAParamCreateNewFrame as QIAP
import logging
#import AnaLyseThematics as AT
#import xlwings

#from difflib_data import *
#from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


# fileName - File name of input document (SDD)
# Returns an array of tables, available in the given document
def is_word_document_corrupted(file_path):
    try:
        docx.Document(file_path)
        return False  # Document loaded successfully, not corrupted
    except Exception as e:
        logging.warning(f"\nFile {os.path.basename(file_path)} is not valid....")
        return True  # Document is corrupted


def getTables(fileName):
    wordDoc = docx.Document(fileName)
    # logging.info("wordDoc.tables - ", wordDoc.tables)
    return wordDoc.tables


def save_as_docx(path):
    # Opening MS Word
    try:
        import win32com.client as win32
        from win32com.client import constants

        word = win32.gencache.EnsureDispatch('Word.Application')


    except AttributeError:
        logging.info("e")
        MODULE_LIST = [m.__name__ for m in sys.modules.values()]
        for module in MODULE_LIST:
            if re.match(r'win32com\\.gen_py\\..+', module):
                del sys.modules[module]
        shutil.rmtree(os.path.join(os.environ.get('LOCALAPPDATA'), 'Temp', 'gen_py'))
        import win32com.client as win32
        from win32com.client import constants
        word = win32.gencache.EnsureDispatch('Word.Application')

    doc = word.Documents.Open(path)
    doc.Activate()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(new_file_abs, FileFormat=constants.wdFormatXMLDocument)
    doc.Close(False)

    logging.info("DOCX File Path", os.path.realpath(new_file_abs))
    return os.path.realpath(new_file_abs)


# tables - list of tables available in SDD
# keyword - signal name
# returns - -1 if none is found else it will return the table which has the signal name
def t_FindTable(table, keyword):
    #logging.info("findinf table")
    if table!=-1:
        if (len(table.columns) < 4):
            for row in table.rows:
                for cell in row.cells:
                    if len(cell.text) !=0:
                        if keyword in cell.text:
                            if keyword.find("REQ-") != -1:
                                return table
                            else:
                                if cell.text.find("REQ-") == -1:
                                    return table
    return -1


def threading_findTable(tables,keyword):
    futures=[]
    nThreads=len(tables)
    logging.info("no. of threads", nThreads)
    with ThreadPoolExecutor(max_workers=nThreads) as exe:
        for table in tables:
            futures.append(exe.submit(t_FindTable, table, keyword))
         
        for future in concurrent.futures.as_completed(futures):
            if future.result()!=-1:
                logging.info("Table Found")
                for f in futures:
                    f.cancel()
                #concurrent.futures.Future.cancel()
                return future.result()
    return -1


def findTable(tables, keyword):
    logging.info("KEYWORDS----->>>>", keyword)
    logging.info("findinf table")
    outTable=-1
    if tables!=-1:
        logging.info("YESS")
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    if len(cell.text) !=0:
                        # logging.info("Cell Text", "*"+keyword+"*")
                        if keyword in cell.text:
                            logging.info("IN IF CONDITION")
                            outTable=table
    # logging.info("findinf table",len(outTable.columns),len(outTable.rows))
    return outTable


# Takes Table and Keyword as argument. This function will search the table with given keyword.
# Returns a cell if match is found, If no match found it returns -1
def searchTable(table, keyword):
    text = ''
    if table != -1:
        i=0
        for col in table.columns:
            j=0
            for cell in col.cells:
                if len(cell.text) !=0:
                    if keyword in cell.text:
                        text=table.columns[i+1].cells[j].text
                j=j+1
            i=i+1
    else:
        text=-1
    return text


def getOldContents(table,keyword):
    text = ''
    logging.info("In getOldContents function", keyword)
    if table != -1:
        i=0
        for col in table.columns:
            j=0
            for cell in col.cells:
                clearStrikethrough(cell)
                if len(cell.text) !=0:
                    if keyword in cell.text:
                        # logging.info("keyword in getOldContents = ", keyword)
                        try:
                            clearStrikethrough(table.columns[i+1].cells[j])
                            text=(table.columns[i+1].cells[j].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            # logging.info("text1", text)
                        except Exception as e:
                            logging.info("Exception in getOldContents = ", e)
                            text = (table.columns[i].cells[j + 1].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            # logging.info("text2", text)
                j=j+1
            i=i+1
    else:
        text=-1
    logging.info("In getOldContents function returning text", text)
    return text


def getNewContents(table, keyword):
    logging.info("In getNewContents function finding cnt")
    text = -1
    data = {'req': keyword}
    if table != -1:
        i = 0
        for col in table.columns:
            j = 0
            for cell in col.cells:

                if len(cell.text) != 0:
                    # logging.info("t",cell.text)
                    if "n°requirement" in cell.text.lower():
                        clearStrikethrough(table.columns[i].cells[j + 1])
                        text = (table.columns[i].cells[j + 1].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        oldName = re.findall(r'GEN-.*\)', text)
                        newName = re.findall(r'REQ-.*', text)
                        data["oldName"] = oldName
                        data['req'] = newName
                    if "content of the requirement" in cell.text.lower():
                        clearStrikethrough(table.columns[i].cells[j + 1])
                        text = (table.columns[i].cells[j + 1].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["content"] = text
                        # logging.info("new text2",text)
                    elif "effectivity expression" in cell.text.lower():
                        clearStrikethrough(table.columns[i].cells[j + 1])
                        text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["effectivity"] = text
                    elif "lcdv" in cell.text.lower():
                        clearStrikethrough(table.columns[i].cells[j + 1])
                        text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["lcdv"] = text
                    elif "diversity" in cell.text.lower():
                        clearStrikethrough(table.columns[i].cells[j + 1])
                        text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["diversity"] = text
                    elif "target configuration" in cell.text.lower():
                        clearStrikethrough(table.columns[i].cells[j + 1])
                        text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["target configuration"] = text
                j = j + 1
            i = i + 1
    else:
        data = -1
    # logging.info("In getNewContents function returning text",text)
    return data


def clearStrikethrough(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # Check whether the run text is strikethrough
            if run.font.strike:
                # Remove the strikethrough text
                run.clear()


def getThematics(table, keyword):
    text = ''
    logging.info("In getThematics function findinf thm")
    if table != -1:
        i=0
        for col in table.columns:
            j=0
            for cell in col.cells:
                if len(cell.text) !=0:
                    # logging.info("t------>>>", cell.text)
                    if "LCDV" in cell.text or "Diversity Expression" in cell.text or "Effectivity Expression" in cell.text or "TARGET CONFIGURATION" in cell.text:
                        try:
                            logging.info("value of i & j---->", i, j)
                            text = str(table.columns[i+1].cells[j].text.encode('utf-8').strip())
                            logging.info("Text----->", text)
                            # text = text.decode('utf-8').strip()
                            # logging.info("type = ", type(text))
                            # logging.info("In getThematics function text1 - ",text, "\n" + str(text))
                        except Exception as e:
                            logging.info("Exception in getThematics = ", e)
                            text = str(table.columns[i].cells[j+1].text.encode('utf-8').strip())
                            # text = text.decode('utf-8').strip()
                            # logging.info("In getThematics function text2 - ",text)
                j = j+1
            i=i+1
    else:
        text=-1
    logging.info("In getThematics function returning text",text)
    return text


def getOldThematics(table,keyword):
    text = ''
    if table != -1:
        i=0
        for col in table.columns:
            j=0
            for cell in col.cells:
                if len(cell.text) !=0:
                    if keyword in cell.text:
                        text=table.columns[i].cells[j].text
                j=j+1
            i=i+1
    else:
        text=-1
    return text


def oldToNewThm(listOfContents):
   #keys=['=','AND','OR']  
   brackets = ['(',')','[',']','{','}']
   newThm = []
   found = 0
   last_index = -1 #to get index of duplicate elements#
   referntialEC = EI.openReferentialEC()
   logging.info("referential EC opened in oldToNewThm",listOfContents)
   for w in listOfContents:
       if w == '=':
           if re.search("_",listOfContents[listOfContents.index(w)-1]).group():
               thm = listOfContents[listOfContents.index(w, last_index+1)-1]
               val = listOfContents[listOfContents.index(w, last_index+1)+1]
               logging.info("thm & val before = ", thm, val)
               if last_index < (len(listOfContents)-1):
                   last_index = listOfContents.index(w, last_index+1)
               else:
                   last_index = -1
               thmCopy = thm
               valCopy = val
               for c in thm:
                   if c in brackets:
                       found=1
                       bIndex=thm.index(c)
                       b=c
                       thmCopy=thmCopy.replace(c,'') 
               for c in val:
                   if c in brackets:
                       found=1
                       bIndex=val.index(c)
                       b=c
                       valCopy = valCopy.replace(c, '')
               logging.info("thm & val after = ", thmCopy, valCopy)
               newThmName=EI.getNewThematics(thmCopy, valCopy, referntialEC)
               logging.info("newThmName = ", newThmName)
               if newThmName != -1:
                   if found==1:
                       found=0
                       if bIndex==0:
                           newThmName=b+newThmName
                       else:
                           newThmName=newThmName+b
                   newThm.append(newThmName)
               else:
                   logging.info("returning -1 from oldToNewThm")
                   referntialEC.close()
                   logging.info("referential EC closed in oldToNewThm and returning -1")
                   return -1
       if w =='AND':
           newThm.append(w)
       if w=='OR':
           newThm.append(w)
   logging.info("returning newThm from oldToNewThm")
   referntialEC.close()
   logging.info("referential EC closed in oldToNewThm", newThm)
   if len(newThm) != 0:
       return newThm
   else:
       return -1


def checkFormat(table, keyword):
    logging.info("In checkFormat functiom")
    networkList = ["HS1","HS2","HS3","HS4","HS5","HS6","HS7","FD3","FD7","FD8","LIN_VSM_1","LIN_VSM_2","LIN_VSM_3","LIN_VSM_4","LIN_VSM_5","LIN_VSM_6","LIN_VSM_7"]
    keywordList = ["synchronous frame gateway","frame gateway","asynchronous signal gateway","synchronous gateway"]
    old = 0
    for i, rows in enumerate(table.rows):
        for j, cell in enumerate(rows.cells):
            # logging.info("d--->", cell.text)
            if ("LCDV" in cell.text) or ("content of the requirement" in cell.text.lower()):
                # logging.info("New Table Format")
                old=1
            elif keyword in cell.text:
                # logging.info("data in gateway ------> ", table.rows[i].cells[j+1].text)
                if ("gateway" in table.rows[i].cells[j].text.lower()) and ([ele for ele in networkList if(ele in table.rows[i].cells[j+1].text)]):
                    old = 2
                    logging.info("Gateway requirement(1)", cell.text)
                elif ("gateway" in table.rows[i].cells[j+1].text.lower()) and ([ele for ele in networkList if(ele in table.rows[i].cells[j+1].text)]):
                    old = 2
                    logging.info("Gateway requirement(11)", cell.text)
                elif ([ele for ele in keywordList if(ele in table.rows[i].cells[j+1].text.lower())]):
                    old = 2
                    logging.info("Gateway requirement(2)", cell.text)
                elif table.rows[i].cells[j + 1].tables:
                    old = 2
                    logging.info("Gateway requirement(3)", cell.text)
    logging.info("Returning old = ", old)
    return old


def loadGWKeywords():#tested ok
    f_keys= open('../user_input/GateWayKeywords.json', "r",encoding='utf8')
    keywords=json.load(f_keys)
    return keywords


def convertGateWayData(keywords,data):#tested
    copyOfData=data.copy()
    for key in keywords["GatewayKeywords"]:
        for text in keywords["GatewayKeywords"][key]:
            #logging.info("t",text,key)
            for dataKey in copyOfData:
                #logging.info("d",dataKey,text)
                if  re.search(text,dataKey):
                    #logging.info("r",dataKey,key)
                    try:
                        data[key]=data.pop(dataKey)
                        #logging.info("after pop",data)
                    except Exception as e:
                        logging.info("e",e)
                        pass
    return data


def convertKeys(keys):
    #ipKeywords=["UpStream","DownStream"]
    listOfData=list(keys)
    logging.info("listOfData = ", listOfData, len(listOfData))
    copyOfData=listOfData.copy()
    flagUp = 0
    flagDown = 0
    for n,k in enumerate(copyOfData):
        if (flagUp == 0):
            if re.search("Upstream".lower(),k.lower()):
                flagUp = 1
                j=n+1
                i=0
                while i<3:
                    logging.info("--++", i, j)
                    listOfData[j]="Upstream".join(listOfData[j].split("/"))
                    j=j+1
                    i=i+1
        if (flagDown == 0):
            if re.search("Downstream".lower(),k.lower()):
                flagDown = 1
                j=n+1
                i=0
                while i<3:
                    logging.info("--", i, j)
                    listOfData[j]="Downstream".join(listOfData[j].split("/"))
                    j=j+1
                    i=i+1
    #logging.info("kk",listOfData)
    return listOfData


def getGateData(tables):
    for table in tables:
        # data = []
        keys = None
        for i, col in enumerate(table.columns):
            text = (cell.text for cell in col.cells)
            # for t in text:
            # logging.info("rxr",t)
            # logging.info("text,",(t for t in text))
            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                text = convertKeys(text)
                keys = tuple(text)
                continue
            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            # logging.info("r",row_data)
            for key in row_data:
                if key == '':
                    for j, cell in enumerate(col.cells):
                        if cell.tables:
                            return getGateData(cell.tables)
            return row_data
            # data.append(row_data)


def parseGateWayData(data):
    result={"UpStreamNetwork":"",
            "UpStreamFrame":"",
            "UpStreamSignal":"",
            "DownStreamNetwork":"",
            "DownStreamFrame":"",
            "DownStreamSignal":"",
            "TEMPO":"",
            "TEMPMIN":"",
            "TEMPMAX":""}
    for key in result:
        for dataKey in data:
            if key==dataKey:
                result[key]=data[dataKey]
    return result


def getGatewayContent(table,keyword):
    logging.info("in gateway contents ",keyword)
    data=-1
    keywords=loadGWKeywords()
    if table != -1:
        #i=0
        for i,col in enumerate(table.columns):
            #j=0
            for j,cell in enumerate(col.cells):
                #logging.info("t ",cell.text)
                if len(cell.text) !=0:
                    if keyword in cell.text:
                        logging.info("row cell ",i,j)
                        # logging.info("keyword in getOldContents = ", keyword)
                        
                        if table.columns[i+1].cells[j].tables:
                            data=getGateData(table.columns[i+1].cells[j].tables)
                            logging.info("data(1)",data)
                            data=convertGateWayData(keywords,data)
                            logging.info("data(2)", data)
                            data=parseGateWayData(data)
                            logging.info("data(3)",data)
    else:
        data=-1
    logging.info("In getOldContents function returning text", data)
    return data


def getThematicsGateway(table, keyword):
    text = ''
    logging.info("In getThematics function findinf thm")
    if table != -1:
        i = 0
        for col in table.columns:
            j = 0
            for cell in col.cells:
                if len(cell.text) != 0:
                    if keyword in cell.text:
                        try:
                            if "TARGET CONFIGURATION" in table.columns[i].cells[j+1].text:
                                try:
                                    text = str(table.columns[i + 1].cells[j+1].text.encode('utf-8').strip())
                                except Exception as e:
                                    logging.info("Exception in getThematicsGateway = ", e)
                                    text = -1
                            else:
                                text = -1
                        except IndexError:
                            logging.info("Exception in getThematicsGateway(2) - maybe due to req is present in last row and no thematique found")
                            text = -1
                j = j + 1
            i = i + 1
    else:
        text = -1
    logging.info("In getThematics function returning text", text)
    return text


def findTableOfContent(tables, keyword):
    logging.info("KEYWORDS----->>>>", keyword)
    logging.info("findinf table")
    outTable = -1
    contentOfReq = []
    isOnlyTable = 0
    if tables != -1:
        logging.info("YESS")
        for table in tables:
            # logging.info("table.rows ", table.rows)
            flowReq = ""
            i = 0
            for row in table.rows:
                # logging.info("\n\nrow.... ", row.cells[0].text)
                for cell in row.cells:
                    if len(cell.text) != 0:
                        # find_req = re.findall(r'\b(REQ|GEN|VHL)\-\S+', cell.text)
                        if cell.text.startswith('REQ') or cell.text.startswith('GEN') or cell.text.startswith('VHL'):
                            find_req = cell.text.split("\n")
                            # logging.info(f"find_req111 {find_req}")
                            # logging.info(f"cell.text {cell.text}")
                            if find_req:
                                flowReq = find_req[0]
                        if "content of the requirement" in cell.text.lower():
                            i = 1
                            continue
                        # if keyword in cell.text and i == 1:
                        if i == 1:
                            extracted_strings = extract_string_with_underscore(cell.text)
                            if extracted_strings:
                                for ex_str in extracted_strings:
                                    if keyword.strip() == ex_str.strip():
                                        logging.info(f"cell.text {cell.text}")
                                        logging.info(f"IN IF CONDITION {cell.text}")
                                        isOnlyTable = 0
                                        outTable = table
                                        if cell.text not in contentOfReq:
                                            contentOfReq.append(f"{flowReq}==> {cell.text}")
                                        i = 0
                        elif keyword in cell.text and i != 1:
                        # elif keyword == cell.text and i != 1:
                            logging.info(f"IN IF CONDITION1 {cell.text}")
                            isOnlyTable = 1
                            logging.info("isOnlyTable ...+++++ ", isOnlyTable)

    logging.info("isOnlyTable ... ", isOnlyTable)
    # logging.info("findinf table",len(outTable.columns),len(outTable.rows))
    return contentOfReq, isOnlyTable


def getContent(Doc,ReqName,ReqVer):
    logging.info("Doc,ReqName,ReqVe------->",Doc,ReqName,ReqVer)
    Content = ''
    RqTable = ''
    try:
        TableList=getTables(Doc)
        logging.info("TableList------>",TableList)
        #RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
        RqTable=threading_findTable(TableList, ReqName)
        logging.info("")
        if RqTable==-1:
            if ((ReqName.find('.')!=-1)|(ReqName.find('_')!=-1)):
                ReqName=ReqName.replace('.', '-')
            RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
        else:
            RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
        if RqTable!=-1:
            chkOldFormat=checkFormat(RqTable, ReqName+"("+ReqVer+")")
            if chkOldFormat==0:
                Content=getOldContents(RqTable, ReqName+"("+ReqVer+")")
            else:
                Content=getNewContents(RqTable, ReqName+"("+ReqVer+")")
        else:
            RqTable=threading_findTable(TableList, ReqName+" "+ReqVer)
            if RqTable!=-1:
                chkOldFormat=checkFormat(RqTable, ReqName+" "+ReqVer)
                if chkOldFormat==0:
                    Content=getOldContents(RqTable, ReqName+" "+ReqVer)
                else:
                    Content=getNewContents(RqTable, ReqName+" "+ReqVer)
            else:
                RqTable=threading_findTable(TableList, ReqName+"  "+ReqVer)
                if RqTable!=-1:
                    chkOldFormat=checkFormat(RqTable, ReqName+"  "+ReqVer)
                    if chkOldFormat==0:
                        logging.info("IF ReqName+"  "+ReqVer------->", ReqName+"  "+ReqVer)
                        Content=getOldContents(RqTable, ReqName+"  "+ReqVer)
                    else:
                        logging.info("ELSE ReqName+"  "+ReqVer------->", ReqName + "  " + ReqVer)
                        Content=getNewContents(RqTable, ReqName+"  "+ReqVer)
                else:
                    Content = -1
        logging.info("content123--->",RqTable,Content,Doc)
        req_content_tuples = list(zip(RqTable, Content, Doc))
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        logging.info(f"\nSomething went wrong {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return RqTable, Content, Doc


def oldFormatContentType1(doc, requirement_id):
    for table in doc.tables:
        num_rows = len(table.rows)
        num_columns = len(table.columns)
        if num_columns == 3 and num_rows >= 12:
            req_id_cell = table.rows[0].cells[0]
            req_id = req_id_cell.text.strip()
            if requirement_id in req_id:
                return table
    return -1


def oldFormatContentType2(doc, requirement_id):
    for table in doc.tables:
        num_rows = len(table.rows)
        num_columns = len(table.columns)
        if num_columns == 3 and num_rows <= 6 and table.rows[0].cells[0].text.find("N° Exigence(v)") != -1:
            for row in range(1, num_rows):
                req_id_cell = table.rows[row].cells[0]
                req_id = req_id_cell.text.strip()
                if requirement_id in req_id:
                    return table
    return -1


def getContent_PT(Doc,ReqName,ReqVer):
    logging.info("Doc,ReqName,ReqVe------->",Doc,ReqName,ReqVer)
    Content = ''
    RqTable = ''
    try:
        TableList=getTables(Doc)
        logging.info("TableList------>",TableList)
        #RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
        RqTable=threading_findTable(TableList, ReqName)
        logging.info("")
        if RqTable==-1:
            if ((ReqName.find('.')!=-1)|(ReqName.find('_')!=-1)):
                ReqName=ReqName.replace('.', '-')
            RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
        else:
            RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
        if RqTable!=-1:
            chkOldFormat=checkFormat(RqTable, ReqName+"("+ReqVer+")")
            if chkOldFormat==0:
                Content=getOldContents(RqTable, ReqName+"("+ReqVer+")")
            else:
                Content=getNewContents(RqTable, ReqName+"("+ReqVer+")")
        else:
            RqTable=threading_findTable(TableList, ReqName+" "+ReqVer)
            if RqTable!=-1:
                chkOldFormat=checkFormat(RqTable, ReqName+" "+ReqVer)
                if chkOldFormat==0:
                    Content=getOldContents(RqTable, ReqName+" "+ReqVer)
                else:
                    Content=getNewContents(RqTable, ReqName+" "+ReqVer)
            else:
                RqTable=threading_findTable(TableList, ReqName+"  "+ReqVer)
                if RqTable!=-1:
                    chkOldFormat=checkFormat(RqTable, ReqName+"  "+ReqVer)
                    if chkOldFormat==0:
                        logging.info("IF ReqName+"  "+ReqVer------->", ReqName+"  "+ReqVer)
                        Content=getOldContents(RqTable, ReqName+"  "+ReqVer)
                    else:
                        logging.info("ELSE ReqName+"  "+ReqVer------->", ReqName + "  " + ReqVer)
                        Content = getNewContents(RqTable, ReqName+"  "+ReqVer)
                else:
                    Content = -1
        logging.info("content123--->",RqTable,Content,Doc)
        req_content_tuples = list(zip(RqTable, Content, Doc))
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        logging.info(f"\nSomething went wrong {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return RqTable, Content, Doc


def getThematicContent(Doc, ReqName, ReqVer):
    TableList = getTables(Doc)
    RqTable = threading_findTable(TableList, ReqName)
    # RqTable=threading_findTable(TableList, ReqName)

    if RqTable == -1:
        if ((ReqName.find('.') != -1) | (ReqName.find('_') != -1)):
            ReqName = ReqName.replace('.', '-')
        RqTable = threading_findTable(TableList, ReqName + "(" + ReqVer + ")")

    if RqTable != -1:
        chkOldFormat = checkFormat(RqTable, ReqName + "(" + ReqVer + ")")
        if chkOldFormat == 0:
            return -2
            Content = getOldContents(RqTable, ReqName + "(" + ReqVer + ")")
        else:
            Content = getRawThematic(RqTable, ReqName + "(" + ReqVer + ")")
    else:
        RqTable = threading_findTable(TableList, ReqName + " " + ReqVer)
        if RqTable != -1:
            chkOldFormat = checkFormat(RqTable, ReqName + " " + ReqVer)
            if chkOldFormat == 0:
                return -2
                Content = getOldContents(RqTable, ReqName + " " + ReqVer)
            else:
                Content = getRawThematic(RqTable, ReqName + " " + ReqVer)
        else:
            RqTable = threading_findTable(TableList, ReqName + " " + ReqVer)
            if RqTable != -1:
                chkOldFormat = checkFormat(RqTable, ReqName + " " + ReqVer)
                if chkOldFormat == 0:
                    return -2
                    Content = getOldContents(RqTable, ReqName + "  " + ReqVer)
                else:
                    Content = getRawThematic(RqTable, ReqName + "  " + ReqVer)
            else:
                Content = -1
    return Content


def getRawThematic(table, keyword):
    logging.info("In getNewContents function finding cnt")
    text = -1
    data = {'req': keyword, 'effectivity': "", "lcdv": "", "diversity": ""}
    if table != -1:
        i = 0
        for col in table.columns:
            j = 0
            for cell in col.cells:
                if len(cell.text) != 0:
                    # logging.info("t",cell.text)
                    if "effectivity expression" in cell.text.lower():
                        text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["effectivity"] = text
                    elif "lcdv" in cell.text.lower():
                        text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["lcdv"] = text
                    elif "diversity expression" in cell.text.lower():
                        text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["diversity"] = text
                j = j + 1
            i = i + 1
    else:
        text = -1
    # logging.info("In getNewContents function returning text",text)
    return data


def compareDocs(currDoc, prevDoc, testReqName, testReqVer, oldReq, newReq):
    testReqName = testReqName.strip()
    testReqVer = testReqVer.strip()
    oldReq = oldReq.strip()
    newReq = newReq.strip()
    logging.info("compareDocs parametres = ", testReqName, testReqVer, "oldReq =*"+oldReq+"*", "newReq =*"+newReq+"*")
    if len(newReq) == 0:
        if oldReq.find("(") != -1:
            reqName = oldReq.split("(")[0]
            reqVer = oldReq.split("(")[1].split(")")[0]
        else:
            reqName = oldReq.split(" ")[0]
            reqVer = oldReq.split(" ")[1]
    else:
        if newReq.find("(") != -1:
            reqName = newReq.split("(")[0]
            reqVer = newReq.split("(")[1].split(")")[0]
        else:
            reqName = newReq.split(" ")[0]
            reqVer = newReq.split(" ")[1]
    reqName = reqName.strip()
    reqVer = reqVer.strip()
    logging.info("In compareDocs function", testReqName, testReqVer, "*"+reqName+"*", "*"+reqVer+"*")
    #logging.info("In compareDocs function", testReqName, testReqVer, reqName, reqVer)
    oldContent = getContent(prevDoc,testReqName,testReqVer)
    logging.info("Old content found", oldContent)
    newContent = getContent(currDoc, reqName, reqVer)
    logging.info("New content found", newContent)
    #logging.info(oldRqTable())
    if (type(oldContent==str) and (type(newContent==str))):
    
        logging.info(oldContent)
        logging.info(newContent)
        if ((type(oldContent)==str) and (type(newContent)==str)):
            logging.info("new Content(1) ", newContent.split())
            logging.info("old Content(1) ", oldContent.split())
            for i in oldContent:
                for j in i:
                    if '=' in j:
                        j=j.replace(j," = ")
                oldContent=oldContent.replace(i,j)
            for i in newContent:
                for j in i:
                    if '=' in j:
                        j=j.replace(j," = ")
                newContent=newContent.replace(i,j)

            newContentSplit = newContent.split()
            oldContentSplit = oldContent.split()
            logging.info("new Content(2) ", newContentSplit)
            logging.info("old Content(2) ", oldContentSplit)

            for n, i in enumerate(newContentSplit):
                #logging.info("I = ", i)
                if i == "IF":
                    indexIF = n
                    logging.info(indexIF)
                    del newContentSplit[:indexIF]
                    break
            logging.info("new Content Modified = ", newContentSplit)
            for n, i in enumerate(oldContentSplit):
                if i == "IF":
                    indexIF = n
                    del oldContentSplit[:indexIF]
                    break
            logging.info("old Content Modified = ", oldContentSplit)

            contentDiff = difflib.ndiff(oldContentSplit, newContentSplit)
            diffDict = {"equal": [], "del": [], "add": []}
            
            #create dictionary
            for i in contentDiff:
                if i.startswith('-'):
                    #logging.info("Deleted Word",i.split()[1])
                    diffDict["del"].append(i.split()[1])
                if i.startswith('+'):
                    #logging.info("Added Word",i.split()[1])
                    diffDict["add"].append(i.split()[1])
                if i.startswith(' '):
                    #logging.info("No change",i.split()[0])
                    diffDict["equal"].append(i.split()[0])

            if oldContentSplit == newContentSplit:
                logging.info("No functional Imapct")
                return -1,-1,-1
            else:
                logging.info("In comaoreDocs diffDict = ", diffDict)
                return oldContentSplit,newContentSplit,contentDiff
        else:
            return -2,-2,-2
    else:
        return -2,-2,-2


def extract_string_with_underscore(text):
    # extract the string which present with underscore from the content of requirement
    result = re.findall(r'\w+_\w+', text)
    # logging.info(f"extracted string {result}")

    return result


def getReqContentData(table, keyword):
    logging.info("getReqContentData keyword---------->", keyword)
    logging.info("Getting the thematic....")
    text = -1
    data = {'req': keyword, 'content': "", 'effectivity': "", 'lcdv': "", 'diversity': "", 'target': "", 'celltext':"", 'flow': "",'frame': "", 'flowframe': ""}
    if table != -1:
        i = 0
        for col in table.columns:
            j = 0
            for cell in col.cells:
                # logging.info(f"cell.text {cell.text}")
                if len(cell.text) != 0:

                    if "n°requirement" in cell.text.lower():
                        text = (table.columns[i].cells[j + 1].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        oldName = re.findall(r'GEN-.*\)', text)
                        newName = re.findall(r'REQ-.*', text)
                        data["oldName"] = oldName
                        data['req'] = newName
                    if "content of the requirement" in cell.text.lower():
                        text = (table.columns[i].cells[j + 1].text.encode('utf-8').strip())
                        text = text.decode('utf-8').strip()
                        data["content"] = text
                        logging.info(f"\n\ncell.tables {table.columns[i].cells[j + 1].tables}\n\n")
                        # logging.info("new text2",text)
                    if keyword in cell.text:
                        try:
                            celltext = table.columns[i + 1].cells[j].text
                            logging.info("celltext--->", celltext)
                            rqtable = table.columns[i + 1].cells[j]
                            data['celltext'] = celltext
                            logging.info('text---4444---->', text)
                            content = table.columns[i + 1].cells[j].text
                            # content = (content.split(' ') and content.split('('))
                            logging.info('content---->', content)
                            # data.update({"Content__": content})
                            data['content'] = content
                            # logging.info('data.update({"Content": content})----->', data["Content"])
                            # if Flow
                            b = QIAP.getdocContent(rqtable)
                            logging.info("b--->",b)
                            data["flow"] = b['flows']
                            data["frame"] = b['frame']
                            data["flowframe"] = b['flowframe']
                            data["circuit"] = b['circuit']
                        except:
                            pass
                    if "Comments" in cell.text or "Comment" in cell.text:
                        try:
                            logging.info(f"table.columns[i + 1].cells[j].text {table.columns[i + 1].cells[j].text}")
                            logging.info(f"\n\nuuuuuuuuuuuut {table.columns[i].cells[j-1].text}")
                            if keyword in table.columns[i].cells[j-1].text:
                                logging.info("%^%^%^%^%^%^%^%^%^")
                                text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                                text = text.decode('utf-8').strip()
                                data["comment"] = text
                        except:
                            logging.info("???????????????? --> ", table.columns[i].cells[j+1].text)
                            text = (table.columns[i].cells[j+1].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["comment"] = text
                    if "effectivity expression" in cell.text.lower():
                        try:
                            text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["effectivity"] = text
                        except:
                            text = (table.columns[i].cells[j+1].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["effectivity"] = text
                    if "lcdv" in cell.text.lower():
                        try:
                            text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["lcdv"] = text
                        except:
                            text = (table.columns[i].cells[j+1].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["lcdv"] = text
                    if "diversity expression" in cell.text.lower():
                        try:
                            text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["diversity"] = text
                        except:
                            text = (table.columns[i].cells[j+1].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["diversity"] = text
                    if "target configuration" in cell.text.lower():
                        try:
                            text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["target"] = text
                        except:
                            text = (table.columns[i].cells[j+1].text.encode('utf-8').strip())
                            text = text.decode('utf-8').strip()
                            data["target"] = text
                j = j + 1
            i = i + 1
    else:
        text = -1
    # logging.info("In getNewContents function returning text",text)
    return data


def getReqContent(Doc, req, ver):
    logging.info(f"\nFinding the thematic..... {req,ver}")
    isDocValid = is_word_document_corrupted(Doc)
    RqTable = ''
    if not isDocValid:
        TableList = getTables(Doc)
        logging.info(f"TableList{TableList}\n\n")
        # RqTable = threading_findTable(TableList, req)
        # if RqTable == -1:
        if ((req.find('.') != -1) | (req.find('_') != -1)):
            req = req.replace('.', '-')
        RqTable = threading_findTable(TableList, req + "(" + ver + ")")
        logging.info(f"RqTable123 {RqTable}")
        if RqTable != -1:
            chkOldFormat = checkFormat(RqTable, req + "(" + ver + ")")
            if chkOldFormat == 0:
                logging.info(f"chkOldFormat111 {chkOldFormat}")
                Content = getOldContents(RqTable, req + "(" + ver + ")")
            else:
                Content = getReqContentData(RqTable, req + "(" + ver + ")")
        else:
            RqTable = threading_findTable(TableList, req + " (" + ver + ")")
            logging.info(f"RqTable1234 {RqTable}")
            if RqTable != -1:
                chkOldFormat = checkFormat(RqTable, req + " (" + ver + ")")
                if chkOldFormat == 0:
                    logging.info(f"chkOldFormat111 {chkOldFormat}")
                    Content = getOldContents(RqTable, req + " (" + ver + ")")
                else:
                    Content = getReqContentData(RqTable, req + " (" + ver + ")")
            else:
                RqTable = threading_findTable(TableList, req + " " + ver)
                logging.info(f"RqTable1233333 {RqTable}")
                if RqTable != -1:
                    chkOldFormat = checkFormat(RqTable, req + " " + ver)
                    if chkOldFormat == 0:
                        logging.info(f"chkOldFormat222 {chkOldFormat}")
                        Content = getOldContents(RqTable, req + " " + ver)
                    else:
                        Content = getReqContentData(RqTable, req + " " + ver)
                else:
                    RqTable = threading_findTable(TableList, req + "  " + ver)
                    if RqTable != -1:
                        chkOldFormat = checkFormat(RqTable, req + "  " + ver)
                        logging.info(f"chkOldFormat3333 {chkOldFormat}")
                        if chkOldFormat == 0:
                            Content = getOldContents(RqTable, req + "  " + ver)
                        else:
                            Content = getReqContentData(RqTable, req + "  " + ver)
                    else:
                        Content = -1
    else:
        Content = -1
    # return Content
    return RqTable, Content, Doc



# tpBook=xw.Book(r"C:/Users/10388/Downloads/BSI AUtomation/Input/Tests_20_64_01272_19_00870_FSEE_TURNKEY_CPK4_V5_VSM.xlsm")
# ####*****************INPUT Variables******************
# req="VHL-TF-IHVMDA-0734(2)" #from Analyse, modify format if not same(output of IM003)
# newReqName="REQ-0609619(A)" #from Analyse, modify format if not same(output of IM003)
# vPT="4" #from test sheet obtained from impact tab(output of IM004)
# currVer="6" #from Sommaire AND Analyse
# ipDoc="SSD_HMIF_PARKING_ASSISTANCE_PARK_HMI" #from sommaire AND Analyse
# tpSheet=tpBook.sheets['VSM20_GC_20_64_0025'] #from Impact tab D Column (output of IM004)
#
# currDoc= r"C:\\Users\\9346\\OneDrive - Expleo France\\Desktop\\Team Code\\TeamTesting Documents\\27-12-2021\\[V8]SSD_HMIF_GROUND_LINK_HMI.docx"
# prevDoc=r"C:\\Users\\9346\\OneDrive - Expleo France\\Desktop\\Team Code\\TeamTesting Documents\\27-12-2021\\[V7]SSD_HMIF_GROUND_LINK_HMI.docx"
# reqName="REQ-0632896"
# reqVer="C"
# testReqName="REQ-0632896"
# testReqVer="B"
#
# a,b,c=compareDocs(currDoc,prevDoc,testReqName,testReqVer,reqName,reqVer)
# d,e,f=compareThematics(currDoc,prevDoc,testReqName,testReqVer,reqName,reqVer)
# logging.info("old thm ",a)
# logging.info("new Thm",b)
# logging.info("thm diff ",c)
# logging.info("old content ",d)
# logging.info("new content",e)
# logging.info("content diff ",f)