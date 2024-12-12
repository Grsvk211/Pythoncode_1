import datetime as datetime
date_time = datetime.datetime.now()
import ExcelInterface as EI
import InputConfigParser as ICF
import pygetwindow as pgw
import time
import KeyboardMouseSimulator as KMS
import logging
import threading
import sys
import re
import win32com.client
import os
from docx import Document
import xlwings as xw
import DocumentSearch as DS
import NewRequirementHandler as NRH
import WordDocInterface as WDI
import Backlog_Handler as BH
from collections import Counter


def find_most_common_element(lst):
    # counts = Counter(lst)
    # most_common_element = counts.most_common(1)[0][0]
    # return most_common_element
    counts = Counter(lst)

    # Check if 1 is in the list, if yes, return 1
    if '1' in counts:
        return '1'

    # If 1 is not in the list, return the most common element
    most_common_element = counts.most_common(1)[0][0]
    return most_common_element

def excel_popup(windowName):
    while (True):
        window_title = pgw.getActiveWindowTitle()
        global stop_threads
        excel_windows = pgw.getWindowsWithTitle("Excel")
        for each_excel_window in excel_windows:
            if (windowName.split('.')[0] in each_excel_window.title):
                each_excel_window.minimize()
                each_excel_window.maximize()
                if (each_excel_window.isActive == False):
                    each_excel_window.activate()
                    break
            else:
                each_excel_window.minimize()
        if (pgw.getActiveWindowTitle() == "Microsoft Excel"):
            time.sleep(1)
            KMS.rightArrow()
            time.sleep(1)
            KMS.pressEnter()
        active_window = pgw.getActiveWindow()
        if active_window is not None:
            active_window.minimize()
        if stop_threads:
            break


def main(Name, Start_date):
    projects = []
    Thématiques_inconnues = []
    Thématiques_non_applicables = []
    NT_values = []
    NA_values = []
    print("hi")
    try:
        path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
        print("path ---->", path)
        Campagnec_Book = EI.openExcel(path)
        Campagnec_Book.activate()
        Check_list_sTr_sheet = Campagnec_Book.sheets['Check-list Start']
        sheet_value = Check_list_sTr_sheet.used_range.value
        config_name = EI.searchDataInColCache(sheet_value, 1, 'Config')
        print("config_name--------->", config_name)

        Fun_name14 = EI.searchDataInColCache(sheet_value, 1, 'Writter of the campaign')
        row, col = Fun_name14['cellPositions'][0]
        print(row, col)
        EI.setDataInCell(Check_list_sTr_sheet, (row, col + 4), Name)
        EI.setDataInCell(Check_list_sTr_sheet, (row + 1, col + 4), Start_date)
        data = EI.searchDataInColCache(sheet_value, 4, 'NI')
        data1 = EI.searchDataInColCache(sheet_value, 4, 'Global result')
        print("data1--------->", data1['cellPositions'])
        print("data--------->", data['cellPositions'])
        filtered_data = [item for i, item in enumerate(data['cellPositions']) if i in (1, 2, 3, 6, 7)]
        print(filtered_data)
        for position in filtered_data:
            print('position1-------->', position)
            row, col = position
            EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 1), 'Yes')

        for n, tuple_data in enumerate(config_name['cellPositions'][1:]):
            print(tuple_data)
            row, col = tuple_data
            # row, col = config_name['cellPositions'][1]
            print(row, col)
            # if EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 4)) != 'No' and EI.getDataFromCell(
            #         Check_list_sTr_sheet, (row, col + 4)) != 'NI':
            Type_of_Validation = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 4))
            Configurations = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 1))
            Priority = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 5))
            Name_of_Project = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
            print("DFGSFDAFGAERF------->", Configurations, Type_of_Validation, Priority, Name_of_Project)

            EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 5), n + 1)
            if str(n+1) == str(1):
                EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 4), 'PC1')
                Config_sheet = 'Campagne Config 1'
                print("Config_sheet1-------->", Config_sheet)
                project1= EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                projects.append(project1.split('(')[0].strip())
                Thématiques_inconnue, Thématiques_non_applicable, NT_value, NA_value = getthematics(Config_sheet)
                NT_values.append(NT_value)
                NA_values.append(NA_value)
                Thématiques_inconnues.append(Thématiques_inconnue)
                Thématiques_non_applicables.append(Thématiques_non_applicable)

            else:
                n = n-1
                EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 4), 'Compl. PC'+str(n + 1))
                Config_sheet = 'Campagne Config '+str((n+1)+1)
                print("Config_sheet2-------->",Config_sheet)
                project2 = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                Thématiques_inconnue, Thématiques_non_applicable, NT_value, NA_value = getthematics(Config_sheet)
                NT_values.append(NT_value)
                NA_values.append(NA_value)
                Thématiques_inconnues.append(Thématiques_inconnue)
                Thématiques_non_applicables.append(Thématiques_non_applicable)
                projects.append(project2.split('(')[0].strip())
        print("projects.append(project1)---------->", projects)
        print("1stsheet--------->", Thématiques_inconnues, Thématiques_non_applicables)
        print("NT_values and NA_values--------->", NT_values, NA_values)
        Campagnec_Book.save()

    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values


def getthematics(Config_sheet):
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
    print("path ---->", path)
    Campagnec_Book = EI.openExcel(path)
    Campagnec_Book.activate()
    Campagne_sheet = Campagnec_Book.sheets[Config_sheet]
    sheet_value = Campagne_sheet.used_range.value
    Fun_name4 = EI.searchDataInColCache(sheet_value, 1, 'Function(s) or type of test')
    row, col = Fun_name4['cellPositions'][0]
    time.sleep(2)
    print(row, col)
    Fun_name4_1 = EI.searchDataInColCache(sheet_value, 5, 'Thématiques inconnues')
    row, col1 = Fun_name4_1['cellPositions'][0]
    time.sleep(2)
    print("row, col0000------------------>", row, col1)

    Thématiques_inconnues = EI.getDataFromCell(Campagne_sheet, (row, col1 + 1))
    print("Thématiques_inconnues---->", Thématiques_inconnues)

    Fun_name4_2 = EI.searchDataInColCache(sheet_value, 5, 'Thématiques non applicables')
    row, col2 = Fun_name4_2['cellPositions'][0]
    time.sleep(2)
    print("row, col0000------------------>", row, col2)
    Thématiques_non_applicables = EI.getDataFromCell(Campagne_sheet, (row, col2 + 1))
    print("Thématiques_non_applicables---->", Thématiques_non_applicables)

    NT_values = EI.getDataFromCell(Campagne_sheet, (row, col2 + 3))
    print("NT_values------>",NT_values)
    NA_values = EI.getDataFromCell(Campagne_sheet, (row+1, col2 + 3))
    print("NA_values------>", NA_values)

    NT_value = int(re.search(r':\s*(\d+)', NT_values).group(1)) if re.search(r':\s*(\d+)', NT_values) else None
    NA_value = int(re.search(r':\s*(\d+)', NA_values).group(1)) if re.search(r':\s*(\d+)', NA_values) else None

    print("NT_value------->", NT_value, NA_value)

    return Thématiques_inconnues, Thématiques_non_applicables, NT_value, NA_value


def getrequired_software():
    required_software = ''
    software = 'NEA R1.2 SP3 V6 BR2'
    # Define a regular expression pattern to match "NEA R1.2"
    pattern = r'(NEA R1\.2)'

    # Use re.search to find the pattern in the input string
    match = re.search(pattern, software)

    # Check if the pattern is found
    if match:
        # Extract the matched group
        required_software = match.group(1)
        print(required_software)
    else:
        print("Pattern not found.")
    return required_software


def getrequired_project_ID(projects):
    project_ID = ''
    input_string = projects[0]
    # input_string = "Gene_R1_2_D85_BEV_Manu"
    # Define a regular expression pattern to match a single letter followed by two digits
    pattern = r'([A-Za-z]\d{2}|[A-Za-z]{2}\d{1}|[A-Za-z]{1}\d{1}[A-Za-z]{1}|d{1}[A-Za-z]{1}|d{2}[A-Za-z]{1})'
    # Use re.search to find the pattern in the input string
    match = re.search(pattern, input_string)
    # Check if the pattern is found
    if match:
        # Extract the matched group
        project_ID = match.group(0)
        print(project_ID)
    else:
        print("Pattern not found.")
    return project_ID


def getthematics_in_list(Thématiques_non_applicables):
    # Thématiques_non_applicables = "| ABY_02 | AEM_01 | AFC_04 | AGG_01 | AMD_02 | AMG_01 | APT_04 | AWB_01 | AXG_01 | AXW_02 | AZX_02 | BPR_01 | CAE_04 | CAE_05 | CMP_02 | CRA_01 | CWF_03 | D12_01 | D41_01 | D41_02 | D59_02 | D7A_02 | D7K_01 | D7U_02 | D7U_04 | D7W_00 | D7W_03 | DAO_01 | DE7_03 | DI6_04 | DO8_00 | DRG_40 | DUE_07 | DUE_08 | DVQ_64 | DVQ_67 | DXD_00 | DXD_03 | DXD_05 | ENK_00 | EPD_00 | ICS_07 | IPD_00 | IWA_01 | IWH_00 | IWW_00 | LVM_02 | LXA_01"
    Thématiques_non_applicables = Thématiques_non_applicables.strip('|')
    # Split the string into a list and remove leading/trailing whitespaces from each element
    thematiques_list = [theme.strip() for theme in Thématiques_non_applicables.split('|')]
    # Remove any empty strings from the list
    thematiques_lists = list(filter(None, thematiques_list))
    # Print the list
    print(thematiques_lists)
    return thematiques_lists


def not_thematic(project, Thématiques_non_applicables, Thématiques_inconnues, desired_architecture):
    print("project---------->",project)
    Thématiques_inconnues_after_check = ''
    Thématiques_non_applicable = Thématiques_non_applicables
    print("ko")
    result_tuples = []
    thematiques_lists = getthematics_in_list(Thématiques_non_applicable)
    silhouette_files = [file for file in os.listdir(ICF.getInputFolder()) if file.lower().endswith((".xlsm", ".xlsx")) and "Silhouette" in file]
    matching_files = []
    for file in silhouette_files:
        # Check if any substring in the project matches the file name
        if any(substring in project for substring in file.split('_')):
            matching_files.append(file)
    print("matching_files--------------->", matching_files)
    path = ICF.getInputFolder() + "\\" + matching_files[0]
    print("path ---->", path)
    Silhouette_Book = EI.openExcel(path)
    Silhouette_Book.activate()
    Silhouettes_sheet = Silhouette_Book.sheets['Silhouettes']
    sheet_value = Silhouettes_sheet.used_range.value
    # for project in project:
    print("project---------->", project)
    try:
        for result_string in thematiques_lists:
            print("thematiques_lists result_string------>", result_string)
            cleaned_result_string = result_string.strip()
            Fun_name9 = EI.searchDataInColCache(sheet_value, 1, cleaned_result_string)
            # print("Fun_name9----->",Fun_name9)
            # Check if 'cellPositions' is not empty before accessing its elements
            if Fun_name9['cellPositions']:
                row, col = Fun_name9['cellPositions'][0]
                time.sleep(2)
                # print(row, col)
                Project_cell = EI.searchDataInExcelCache(sheet_value, project)
                # print("project cells----------->", Project_cell)
                row18, col18 = Project_cell['cellPositions'][0]
                project1 = EI.getDataFromCell(Silhouettes_sheet, (row18, col18))
                if project == project1:
                    # print("hiiiiiiiii")
                    Silhouettes_sheet_Valeur = EI.getDataFromCell(Silhouettes_sheet, (row, col18))
                    print("ValeurValeurValeurValeur--->", Silhouettes_sheet_Valeur)
                    if Silhouettes_sheet_Valeur == '--' or Silhouettes_sheet_Valeur == 'X' or Silhouettes_sheet_Valeur == 'opt':
                        print("v--------------->", result_string)
                        result_tuples.append((cleaned_result_string, Silhouettes_sheet_Valeur))
        print("result_tuples---------------->", result_tuples)

    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    #
    Thématiques_sheet = Silhouette_Book.sheets['Thématiques']
    sheet_value1 = Thématiques_sheet.used_range.value
    # for project in project:
    print("project---------->", project)
    result__strings = []
    result_Xstrings = []
    try:
        NEA_R1_cell_Xdatas = []
        NEA_R1_cell__datas = []
        Thématiques_inconnuess = getthematics_in_list(Thématiques_inconnues)
        for result_string in Thématiques_inconnuess:
            print("Thématiques_inconnuess result_string------>", result_string)
            cleaned_result_string = result_string.strip()
            Fun_name9 = EI.searchDataInColCache(sheet_value1, 1, cleaned_result_string)
            print("Fun_name9----->", Fun_name9)
            # Check if 'cellPositions' is not empty before accessing its elements
            if Fun_name9['cellPositions']:
                row23, col23 = Fun_name9['cellPositions'][0]
                time.sleep(2)
                print("row23, col23--------->", row23, col23)
                print("desired_architecture----------->", desired_architecture)
                NEA_R1_cell = EI.searchDataInExcelCache(sheet_value1, desired_architecture)
                print("NEA_R1_cell--------->", NEA_R1_cell)
                if NEA_R1_cell['cellPositions']:
                    row15, col15 = NEA_R1_cell['cellPositions'][0]
                    print("row, col15----------->", row15, col15)
                    # NEA_R1_cell_data = EI.getDataFromCell(sheet_value1, (row23, col15))
                    cell_data = EI.getDataFromCell(Thématiques_sheet, (row23, col15))
                    print(f'{result_string} thematics Arch {cell_data}')
                    if cell_data == 'X':
                        NEA_R1_cell_Xdatas.append(cell_data)
                        result_Xstrings.append(result_string)
                    if cell_data == '--':
                        NEA_R1_cell__datas.append(cell_data)
                        result__strings.append(result_string)
            else:
                print(f'Thematics {result_string} not present in the Silhouette file Thematcs sheet. Check in the EC file.')

        print("NEA_R1_cell__datas---------->",NEA_R1_cell__datas, result__strings)
        Thématiques_inconnues_after_check = '|' + '|'.join(result__strings)
        print("Thématiques_inconnues_after_check--------->", Thématiques_inconnues_after_check)

        if result_Xstrings:
            print("NEA_R1_cell_Xdatas---------->", NEA_R1_cell_Xdatas, result_Xstrings)
            print(f'check the thematcis {result_Xstrings} are applicable for the {desired_architecture}. when searching for Thématiques_inconnues.')

    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")

    Silhouette_Book.close()
    return result_tuples, Thématiques_inconnues_after_check


def find_sheet_in_plm_book(PLM_EE_Book, project_ID, required_software):
    for sheet in PLM_EE_Book.sheets:
        sheet_name_parts = sheet.name.split('_')
        if len(sheet_name_parts) == 3:
            sheet_project_ID = sheet_name_parts[1]
            sheet_required_software = sheet_name_parts[2]
            if sheet_project_ID == project_ID and sheet_required_software == required_software:
                return sheet
    return None


# Function to check if a sheet is present in an Excel file
def is_sheet_present(file_path, sheet_name, filename):
    try:
        global stop_threads
        stop_threads = False
        t1 = threading.Thread(target=excel_popup, args=(filename,))
        t1.start()
        # Open the Excel file
        workbook = xw.Book(file_path)
        stop_threads = True

        # Check if the sheet is present
        if any(sheet_name.strip().lower() == sheet.name.strip().lower() for sheet in workbook.sheets):
            return True
        else:
            workbook.close()
            return False
    except Exception as e:
        print(f"Error: {e}")
        return False


def to_find_sheet_inPLM(desired_sheet_name):
    sheet_found_flag = 0
    file_name = ''
    input_folder = ICF.getInputFolder()
    # List of file extensions to filter Excel files
    file_extensions = [".xlsx", ".xlsm"]
    # Sheet name to check for
    # desired_sheet_name = "ENV_D85_NEA R1.2"
    # desired_sheet_name = "VIJAYA345678"

    # Flag to track whether the sheet is found in any file
    sheet_found = False

    # Loop through each file in the folder
    for file_name in os.listdir(input_folder):
        # if any(file_name.endswith(ext) and "PLM" in file_name for ext in file_extensions):
        if any(file_name.endswith(ext) for ext in file_extensions):
            # if file_name.endswith(file_extension) and "PLM" in file_name:
            file_path = os.path.join(input_folder, file_name)

            # Check if the desired sheet is present in the current Excel file
            if is_sheet_present(file_path, desired_sheet_name,file_name):
                sheet_found_flag = 1
                print(f"The sheet '{desired_sheet_name}' is present in the file: {file_name}")
                # Set the flag to True and break out of the loop
                sheet_found = True
                break
    # Check if the sheet was not found in any file
    if not sheet_found:
        print(f"The sheet '{desired_sheet_name}' is not present in any PLM files.")
    return sheet_found_flag, file_name


def checkinplm(result_tuples, file_name, NT_value, NA_value, desired_sheet_name, sheet_found_flag, Thématiques_inconnues, desired_architecture):
    print("desired_sheet_name--------------->", desired_sheet_name)
    # version = 'NEA R1'
    NA_content = []
    # path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[22]
    path = ICF.getInputFolder() + "\\" + file_name
    print("path ---->", path)
    global stop_threads
    stop_threads = False
    # t1 = threading.Thread(target=excel_popup, args=(EI.findInputFiles()[22],))
    t1 = threading.Thread(target=excel_popup, args=(file_name,))
    t1.start()
    PLM_EE_Book = EI.openExcel(path)
    stop_threads = True
    PLM_EE_Book.activate()
    try:
        if sheet_found_flag:
            func_sheet = PLM_EE_Book.sheets[desired_sheet_name]
            # Convert the list to an Excel range object
            sheet_range = func_sheet.range('A1').expand()
            # Find the last non-empty row
            maxrow = sheet_range.end('up').row
            print("maxrow--------->", maxrow)
            sheet_value = func_sheet.used_range.value
            Fun_name9 = EI.searchDataInColCache(sheet_value, 7, 'Name')
            print("Fun_name9----->", Fun_name9)
            # Check if 'cellPositions' is not empty before accessing its elements
            if Fun_name9['cellPositions']:
                row, col = Fun_name9['cellPositions'][0]
                time.sleep(2)
                print(row, col)
                for i in result_tuples:
                    if i[1] == '--' or i[1] == 'X':
                        print("i--------->", i[0])
                        result_tuple = i[0]
                        print("result_tuple.split('_')[0]---->", result_tuple.split('_')[0])
                        try:
                            PLM_sheet_Valeur = EI.searchDataInColCache(sheet_value, col, result_tuple.split('_')[0])
                            print("PLM_sheet_Valeur--------->", PLM_sheet_Valeur)
                            # Condition 0
                            if PLM_sheet_Valeur['count'] == 0:
                                print("PLM_sheet_Valeur['count'] == 0")
                                b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA'
                                NA_content.append(b)

                            elif PLM_sheet_Valeur['count'] == 1:
                                print("PLM_sheet_Valeur['count'] == 1")
                                row, col = PLM_sheet_Valeur['cellPositions'][0]
                                thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                # Condition 1
                                if i[1] == '--':
                                    #  -- !=
                                    if result_tuple != thematic_value:
                                        b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is fixed on the value ' + thematic_value
                                        NA_content.append(b)
                                    else:
                                        b = result_tuple + ':  Raise QIA of Silhouette.'
                                        NA_content.append(b)
                                # Condition 2
                                if i[1] == 'X':
                                    if result_tuple != thematic_value or result_tuple == thematic_value:
                                        b = result_tuple + ':  Raise QIA of Silhouette.'
                                        NA_content.append(b)

                            elif PLM_sheet_Valeur['count'] >= 2:
                                print("PLM_sheet_Valeur['count'] > 2")
                                result_list = []
                                for cellposition in PLM_sheet_Valeur['cellPositions']:
                                    row, col = cellposition
                                    print("row,col", row, col)
                                    thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                    b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is fixed on the value ' + thematic_value
                                    NA_content.append(b)
                                    # Break the loop after processing the first cell position
                                    break
                                #     result_list.append(thematic_value)
                                # result = "|".join(result_list)
                                # b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is fixed on the value ' + result
                                # NA_content.append(b)

                        except Exception as ex:
                            exc_type, exc_obj, exc_tb = sys.exc_info()
                            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                            print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
                            print("Thematics are not present in the PLM_Sheet.")
                    if i[1] == 'opt':
                        print("i--------->", i[0])
                        result_tuple = i[0]
                        print("result_tuple.split('_')[0]---->", result_tuple.split('_')[0])
                        try:
                            PLM_sheet_Valeur = EI.searchDataInColCache(sheet_value, col, result_tuple.split('_')[0])
                            print("PLM_sheet_Valeur--------->", PLM_sheet_Valeur)
                            if PLM_sheet_Valeur['count'] > 0:
                                print("PLM_sheet_Valeur['count'] > 0")
                                result_list = []
                                for cellposition in PLM_sheet_Valeur['cellPositions']:
                                    row, col = cellposition
                                    print("row,col", row, col)
                                    thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                    result_list.append(thematic_value)
                                    b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is applicable'
                                    NA_content.append(b)
                                # result = "|".join(result_list)
                                # b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is applicable'
                                # # b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is optional on the value ' + result
                                # NA_content.append(b)
                        except Exception as ex:
                            exc_type, exc_obj, exc_tb = sys.exc_info()
                            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                            print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
                            print("Thematics are not present in the PLM_Sheet.")

    except:
        print("sheet are not present in the PLM_Sheet.")


    # project_text_bold = "\033[1m" + projects[0] + "\033[0m"
    # print(project_text_bold)
    # doc_content = [projects[0] + ':', '\n ', '\n ', 'Total number of NT files: 178','Total number of NA files: 285', '\n ', ' \n', 'NA because of following thematiques: ', '\n ', '\n']
    doc_content = ['Total number of NT files: ' + str(NT_value), 'Total number of NA files: ' + str(NA_value),
                   'NA because of following thematiques: ']
    if desired_architecture == 'NEA R1.x':
        version = 'NEA R1'
        Thématiques_inconnues_COMMENT = ['The following  thematiques:'+Thématiques_inconnues+'are only applicable for '+ version+' according to the CONFIG_THEMATIQUES file.']
    else:
        version = 'NEA R1.x'
        Thématiques_inconnues_COMMENT = [
            'The following  thematiques:' + Thématiques_inconnues + 'are only applicable for ' + version + ' according to the CONFIG_THEMATIQUES file.']

    NA_NT_content = doc_content + NA_content + Thématiques_inconnues_COMMENT
    # NA_NT_content = NA_content
    print("NA_NT_content-------->", NA_NT_content)
    PLM_EE_Book.close()

    return NA_NT_content


def append_to_word_document(contents, output_path):
    doc = Document()

    # Iterate through the content
    for content in contents:
        # Assume the first element is the heading
        heading = content[0]

        # Check if the heading is a list of strings
        if isinstance(heading, list):
            # Join the list of strings into a single string
            heading = ' '.join(heading)

        # Add the heading to the document with bold formatting
        doc.add_paragraph(heading.strip(), style='Heading1')

        # Iterate through the rest of the content
        for line in content[1:]:
            # Treat '\n' as a new line in the Word document
            if line.strip() == '\n':
                doc.add_paragraph()
            else:
                sanitized_line = ''.join(c for c in line if c.isprintable())

                # Check if the line is not empty after sanitization
                if sanitized_line.strip():
                    doc.add_paragraph(sanitized_line.strip())

    doc.save(output_path)


def updateFeps(projects, Fepss):
# def updateFeps():
#     project = 'R1_2_D85_Gene_Manu'
#     Fepss = ['FEPS_117176','FEPS_117177','FEPS_117218','FEPS_119754','FELP_120108','FEPS_119753','FELP_120104','FEPS_119956','FEPS_119954']
    inconnues_list_flag, non_applicables_list_flag = '', ''
    project = projects
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
    print("path ---->", path)
    Campagnec_Book = EI.openExcel(path)
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[1]
    print("path ---->", path)
    test_Book = EI.openExcel(path)
    contents = []
    Feps_contents = []
    for Feps in Fepss:
        numeric_part = ''.join(filter(str.isdigit, Feps))
        print("numeric_part--------->",numeric_part)
        NA_NT_values = []
        try:
            test_Book.activate()
            Feps_sheet = test_Book.sheets['FEPS History']
            sheet_value = Feps_sheet.used_range.value
            Fun_name4 = EI.searchDataInColCache(sheet_value, 3, numeric_part.strip())
            row, col = Fun_name4['cellPositions'][0]
            time.sleep(2)
            print(row, col)
            Impacted_sheets = EI.getDataFromCell(Feps_sheet, (row, col-1))
            Campagnec_Book.activate()
            Campagne_sheet = Campagnec_Book.sheets['Synthèse des campagnes']
            sheet_value = Campagne_sheet.used_range.value
            # Assuming Impacted_sheets is a list of sheet names
            print("Impacted_sheets.split()--------->", Impacted_sheets.split('\n'))
            impacetdSheets = Impacted_sheets.split('\n')
            existing_values = []
            for sheet_name in impacetdSheets:
                if sheet_name:
                    print(f'Do something with sheet: {sheet_name}')
                    Fun_name4 = EI.searchDataInColCache(sheet_value, 1, sheet_name)
                    print("Fun_name4---------->", Fun_name4)
                    # Extract only the cell values that exactly match the target string
                    # filtered_cell_values = [value for value in Fun_name4.get('cellValue', []) if value.startswith(sheet_name + '_')]
                    # filtered_cell_values = [value for value in Fun_name4.get('cellValue', []) if value.startswith(sheet_name + '_') or value.startswith(sheet_name)]
                    filtered_cell_values = [value for value in Fun_name4.get('cellValue', []) if value == sheet_name or value.startswith(sheet_name + '_')]


                    # Create a new dictionary with the same 'cellPositions' and the filtered 'cellValue'
                    filtered_dict = {'count': len(filtered_cell_values), 'cellPositions': [pos for pos, cell_value in zip(Fun_name4['cellPositions'], Fun_name4['cellValue'])
                                                                                           if cell_value in filtered_cell_values], 'cellValue': filtered_cell_values}

                    # Print the result
                    print("filtered_dict-->", filtered_dict)

                    if filtered_dict['cellPositions']:
                        for cellPositions in filtered_dict['cellPositions']:
                            row, col5 = cellPositions
                            print("row, col5----------->", row, col5)
                            existing_value = EI.getDataFromCell(Campagne_sheet, (row, col5+12))
                            print("existing_value---------->", existing_value)
                            # EI.setDataFromCell(Campagne_sheet, (row, col5 + 12), Feps)

                            # HERE I NEED TO IMPLEMENT THE FEPS LOGIC
                            if existing_value:
                                if existing_value is not None:
                                    if Feps not in existing_values:
                                        existing_value += "\n"  # Add a new line if there is an existing value
                                        new_value = existing_value + Feps
                                        print("existing_value if condition-------->", new_value)
                                        lines = new_value.split("\n")

                                        unique_lines = set(lines)
                                        output_string = "\n".join(unique_lines)
                                        print("output_string if condition- set ------->", output_string)

                                        EI.setDataFromCell(Campagne_sheet, (row, col5+12), output_string)
                                    else:
                                        print(
                                            f"FEPS value '{Feps}' is a repetition and not appended to cell.")
                            else:
                                EI.setDataFromCell(Campagne_sheet, (row, col5+12), Feps)
                                print("existing_value else condition-------->", existing_value)
                            print("project0----------->", project)
                            project1 = EI.searchDataInExcelCache(sheet_value, project)
                            # print("project1---------->", project1)
                            if project1['cellPositions']:
                                row15, col15 = project1['cellPositions'][1]
                                print("row, col15----------->", row15, col15)
                                print("project1['cellValue'][1]----------->", project1['cellPositions'][1], project1['cellValue'][1])
                                if project == project1['cellValue'][1]:
                                    project_name = project1['cellValue'][1]
                                    print(f'hiiii {project_name}')
                                    NA_NT_value = EI.getDataFromCell(Campagne_sheet, (row, col15))
                                    print("NA_NT_values------->",NA_NT_value)
                                    NA_NT_values.append(NA_NT_value)
                    else:
                        print(f'Deleted the sheet impacted {sheet_name} by FEPS.')
            print("NA_NT_values-->", NA_NT_values)
            count_NT = NA_NT_values.count('NT')
            count_NA = NA_NT_values.count('NA')
            content = 'For '+Feps+' will be '+ str(count_NT) +' NT and '+str(count_NA)+' NA'
            contents.append(content)

        except Exception as ex:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
            print("FEps are not present in the testplan FEPS History sheet.")

    project_text_bold = "\033[1m" + project + "\033[0m"
    print(project_text_bold)
    doc_content = [project+':']
    NA_NT_content = doc_content + contents
    print("contentscontentscontents---------->", NA_NT_content)
    Campagnec_Book.save()
    Campagnec_Book.close()
    test_Book.close()
    return NA_NT_content


if __name__ == '__main__':
    ICF.loadConfig()
    print("Tool Started")
    start = time.time()
    # FEPS_117176,FEPS_117177,FEPS_117218,FEPS_119754,FELP_120108,FEPS_119753,FELP_120104,FEPS_119956,FEPS_119954
    Name = "RAMA KRISHNA REDDY KOPPULA (RRE - EXP)"
    Software = "SP3_V6.00 Official"
    Start_date = "12/28/2023"
    Fepss = input("Enter the GAELE_reference by giving , : ")
    Fepss = Fepss.split(',')
    # Prompt the user to enter the number of projects
    num_projects = int(input("Enter the number of projects: "))

    # Initialize empty lists to store the desired sheets and architectures
    desired_sheets = []
    desired_architectures = []

    # Loop to get input for each project
    for i in range(num_projects):
        sheet_name = input(f"Enter the desired sheet name for project {i + 1}: ")
        architecture = input(f"Enter the desired architecture for project {i + 1}: ")

        # Check if the entered architecture starts with "NEA R1."
        if architecture.startswith("NEA R1."):
            architecture = "NEA R1.x"

        desired_sheets.append(sheet_name)
        desired_architectures.append(architecture)

    # Print the list of desired sheets and architectures
    print("Desired sheets:", desired_sheets)
    print("Desired architectures:", desired_architectures)

    All_contents = []
    projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values = main(Name, Start_date)
    # Iterate over the lists simultaneously using zip
    for project, ti, tna, nt, na, desired_sheet_name, desired_architecture in zip(projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values, desired_sheets, desired_architectures):
        no_of_NA_NT_contents = updateFeps(project, Fepss)
        result_tuples, Thématiques_inconnues_after_check = not_thematic(project, tna, ti, desired_architecture)
    #     # result_tuples = not_thematic()
        sheet_found_flag, file_name = to_find_sheet_inPLM(desired_sheet_name)
        NA_content = checkinplm(result_tuples, file_name, nt, na, desired_sheet_name, sheet_found_flag, Thématiques_inconnues_after_check, desired_architecture)
        ALl_content = no_of_NA_NT_contents + NA_content
        All_contents.append(ALl_content)
        print("ALl_content-------->", ALl_content)
    print("ALl_content--------------->", All_contents)
    # # Example usage
    # Get the input folder path using ICF.getInputFolder()
    input_folder = ICF.getInputFolder()
    # Create the full output path by joining the input folder with the document name
    output_path = os.path.join(input_folder, 'Output_document.docx')

    append_to_word_document(All_contents, output_path)
    print(f"Word document saved to: {output_path}")
    end1 = time.time()
    print("\nexecution time " + str(end1 - start))
    print("Task Fully Completed")


# FEPS_119939,FEPS_124324
# FEPS_119939,FEPS_119937,FEPS_121345,FEPS_124347,FEPS_124324

# FEPS_121345
# ENV_D85_NEA R1.2
#
# NEA R1.2
# ENV_J4U_NEA R1.2



