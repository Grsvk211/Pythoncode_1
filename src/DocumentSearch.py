import os
import sys

import docx
import time
import re
import logging


# Splits and Returns Requirement ID and Version
def separate_requirement(text):
    if "REQ-" in text:
        pattern = r'REQ-(\d+)\s*(?:\(|)([A-Z])?(?:\)|)'
        match = re.match(pattern, text)
        if match:
            requirement_name = match.group(1)
            requirement_version = match.group(2)
            return "REQ-" + requirement_name, requirement_version
        else:
            return None, None
    elif "GEN-" in text:
        pattern = r'REQ-(\d+)\s*(?:\(|)([A-Z])?(?:\)|)'
        match = re.match(pattern, text)
        if match:
            requirement_name = match.group(1)
            requirement_version = match.group(2)
            return "GEN-" + requirement_name, requirement_version
        else:
            return None, None


# def oldFormatContent(doc, requirement_id):
#     searchResult = {}
#     for table in doc.tables:
#         num_rows = len(table.rows)
#         num_columns = len(table.columns)
#         if num_columns == 3 and num_rows >= 6:
#             req_id_cell = table.rows[0].cells[0]
#             req_id = req_id_cell.text.strip()
#             if requirement_id in req_id:
#                 searchResult.update({"reqId": requirement_id})
#                 for row_index, row in enumerate(table.rows):
#                     try:
#                         if "content of the requirement" in row.cells[0].text.lower().strip():
#                             next_row = table.rows[row_index + 1]
#                             searchResult.update(
#                                 {"content": str(next_row.cells[0].text)})
#                         if "Effectivity" in row.cells[0].text.strip():
#                             next_row = table.rows[row_index + 1]
#                             searchResult.update(
#                                 {"effectivity": str(next_row.cells[0].encode('utf-8').strip())})
#                         elif "LCDV" in row.cells[0].text.strip():
#                             next_row = table.rows[row_index + 1]
#                             searchResult.update(
#                                 {"LCDV": str(next_row.cells[0].text.encode('utf-8').strip())})
#                         elif "diversity" in row.cells[0].text.lower().strip():
#                             next_row = table.rows[row_index + 1]
#                             searchResult.update(
#                                 {"diversity": str(next_row.cells[0].text.encode('utf-8').strip())})
#                         elif "target configuration" in row.cells[0].text.lower().strip():
#                             next_row = table.rows[row_index + 1]
#                             searchResult.update(
#                                 {"target": str(next_row.cells[1].text.encode('utf-8').strip())})
#                     except Exception as exp:
#                         exc_type, exc_obj, exc_tb = sys.exc_info()
#                         exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
#                         logging.info(
#                             f"\nProblem in fetching the content from old format document {exp} line no. {exc_tb.tb_lineno} file name: {exp_fname}********************")
#
#     return searchResult


def oldFormatContent(doc, requirement_id):
    searchResult = {}
    table = ''
    for table in doc.tables:
        num_rows = len(table.rows)
        num_columns = len(table.columns)
        if num_columns == 3 and num_rows >= 6:
            req_id_cell = table.rows[0].cells[0]
            req_id = req_id_cell.text.strip()
            if requirement_id in req_id:
                print("table-----oldoo--->", table)
                searchResult.update({"reqId": requirement_id})
                for row_index, row in enumerate(table.rows):
                    try:
                        if "content of the requirement" in row.cells[0].text.lower().strip():
                            next_row = table.rows[row_index + 1]
                            content_cell = next_row.cells[0]
                            clearStrikethrough(content_cell)
                            searchResult.update({"content": str(content_cell.text)})
                        if "Effectivity" in row.cells[0].text.strip():
                            next_row = table.rows[row_index + 1]
                            effectivity_cell = next_row.cells[0]
                            clearStrikethrough(effectivity_cell)
                            searchResult.update({"effectivity": str(effectivity_cell.encode('utf-8').strip())})
                        elif "LCDV" in row.cells[0].text.strip():
                            next_row = table.rows[row_index + 1]
                            lcdv_cell = next_row.cells[0]
                            clearStrikethrough(lcdv_cell)
                            searchResult.update({"LCDV": str(lcdv_cell.text.encode('utf-8').strip())})
                        elif "diversity" in row.cells[0].text.lower().strip():
                            next_row = table.rows[row_index + 1]
                            diversity_cell = next_row.cells[0]
                            clearStrikethrough(diversity_cell)
                            searchResult.update({"diversity": str(diversity_cell.text.encode('utf-8').strip())})
                        elif "target configuration" in row.cells[0].text.lower().strip():
                            next_row = table.rows[row_index + 1]
                            target_cell = next_row.cells[1]
                            clearStrikethrough(target_cell)
                            searchResult.update({"target": str(target_cell.text.encode('utf-8').strip())})
                    except Exception as exp:
                        exc_type, exc_obj, exc_tb = sys.exc_info()
                        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                        logging.info(
                            f"\nProblem in fetching the content from old format document {exp} line no. {exc_tb.tb_lineno} file name: {exp_fname}********************")
    print("searchResult,table549876-------->", searchResult, table)
    return searchResult, table



# def newFormatContent(doc, requirement_id):
#     searchResult = {}
#     for table in doc.tables:
#         num_rows = len(table.rows)
#         num_columns = len(table.columns)
#         if num_columns == 2 and num_rows >= 5:
#             req_id_cell = table.rows[1].cells[0]
#             req_id = req_id_cell.text.strip()
#
#             if requirement_id in req_id:
#                 try:
#                     searchResult.update({"reqId": requirement_id, "content": clearStrikethrough(table.rows[1].cells[1].text)})
#                 except Exception as e:
#                     logging.info("Exception in getOldContents = ", e)
#                     searchResult.update({"reqId": requirement_id, "content": table.rows[1].cells[1].text})
#                 for row in table.rows:
#                     try:
#                         if "Effectivity" in row.cells[0].text.strip():
#                             searchResult.update(
#                                 {"effectivity": str(row.cells[1].text.encode('utf-8').strip())})
#                         elif "LCDV" in row.cells[0].text.strip():
#                             searchResult.update(
#                                 {"LCDV": str(row.cells[1].text.encode('utf-8').strip())})
#                         elif "diversity" in row.cells[0].text.lower().strip():
#                             searchResult.update(
#                                 {"diversity": str(row.cells[1].text.encode('utf-8').strip())})
#                         elif "target configuration" in row.cells[0].text.lower().strip():
#                             searchResult.update(
#                                 {"target": str(row.cells[1].text.encode('utf-8').strip())})
#                     except Exception as exp:
#                         exc_type, exc_obj, exc_tb = sys.exc_info()
#                         exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
#                         logging.info(
#                             f"\nProblem in fetching the content from new format document {exp} line no. {exc_tb.tb_lineno} file name: {exp_fname}********************")
#
#     return searchResult


def newFormatContent(doc, requirement_id):
    searchResult = {}
    table= ''
    for table in doc.tables:
        num_rows = len(table.rows)
        num_columns = len(table.columns)
        if num_columns == 2 and num_rows >= 5:
            req_id_cell = table.rows[1].cells[0]
            req_id = req_id_cell.text.strip()

            if requirement_id in req_id:
                print("table-----oonew>",table)
                try:
                    content_cell = table.rows[1].cells[1]
                    clearStrikethrough(content_cell)
                    searchResult.update({"reqId": requirement_id, "content": content_cell.text})
                except Exception as e:
                    logging.info("Exception in getOldContents = ", e)
                    searchResult.update({"reqId": requirement_id, "content": table.rows[1].cells[1].text})
                for row in table.rows:
                    try:
                        if "Effectivity" in row.cells[0].text.strip():
                            searchResult.update(
                                {"effectivity": str(row.cells[1].text.encode('utf-8').strip())})
                        elif "LCDV" in row.cells[0].text.strip():
                            searchResult.update(
                                {"LCDV": str(row.cells[1].text.encode('utf-8').strip())})
                        elif "diversity" in row.cells[0].text.lower().strip():
                            searchResult.update(
                                {"diversity": str(row.cells[1].text.encode('utf-8').strip())})
                        elif "target configuration" in row.cells[0].text.lower().strip():
                            searchResult.update(
                                {"target": str(row.cells[1].text.encode('utf-8').strip())})
                    except Exception as exp:
                        exc_type, exc_obj, exc_tb = sys.exc_info()
                        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                        logging.info(
                            f"\nProblem in fetching the content from new format document {exp} line no. {exc_tb.tb_lineno} file name: {exp_fname}********************")
    print("searchResult,table0123549876-------->",searchResult,table)
    return searchResult, table


def find_requirement_content(file_path, requirement_id):
    print(f"finding requirement content {requirement_id}")
    searchResult = {}
    doc = docx.Document(file_path)
    num_tables = len(doc.tables)
    print("Total Number of Tables = ", num_tables)
    req_res, table = newFormatContent(doc, requirement_id)
    print(f"\n>>>>req_res new: {req_res}")
    if req_res:
        searchResult = req_res
    else:
        print("searching data in old format...")
        req_res, table = oldFormatContent(doc, requirement_id)
        print(f"\n>>>>req_res old: {req_res}")
        if req_res:
            searchResult = req_res
    print("searchResult, table, file_pathfind_req---------->",searchResult, table, file_path)
    return searchResult, table, file_path


def clearStrikethrough(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # Check whether the run text is strikethrough
            if run.font.strike:
                # Remove the strikethrough text
                run.clear()


if __name__ == "__main__":

    # file_path = r"C:\Users\vgajula\Downloads\SSVS_SSFD_GEN2_RSP_CVMM_CONTROL_VHL_MISSION_MODES_23Q2.docx"
    # requirement_id = "REQ-0684801 C"
    # file_path = r"C:\Users\vgajula\Downloads\[V1.0][02014_23_00326] DPE-VSM-SSFD-GEN2-CVMMv1_ISS-0211942_v1.docx"
    # requirement_id = "REQ-0664836  A"
    file_path = r"C:\Users\vgajula\Downloads\Identification.docx"
    requirement_id = "REQ-0240231  D"

    req_name, req_version = separate_requirement(requirement_id)

    start_time = time.time()
    requirement_content = find_requirement_content(file_path, req_name)
    end_time = time.time()
    execution_time = end_time - start_time
    logging.info(f"Execution time: {execution_time:.6f} seconds")
    if requirement_content:
        logging.info(f"Requirement ID: {requirement_id}\nContent: {requirement_content}")
    else:
        logging.info(f"Requirement ID '{requirement_id}' not found.")
