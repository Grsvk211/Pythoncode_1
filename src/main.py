import datetime as datetime
date_time = datetime.datetime.now()
import ExcelInterface as EI
import InputConfigParser as ICF
import pygetwindow as pgw
import time
import KeyboardMouseSimulator as KMS
import logging
import threading
import sys
import re
import win32com.client
import os
from docx import Document

def excel_popup(windowName):
    while (True):
        window_title = pgw.getActiveWindowTitle()
        global stop_threads
        excel_windows = pgw.getWindowsWithTitle("Excel")
        for each_excel_window in excel_windows:
            if (windowName.split('.')[0] in each_excel_window.title):
                each_excel_window.minimize()
                each_excel_window.maximize()
                if (each_excel_window.isActive == False):
                    each_excel_window.activate()
                    break
            else:
                each_excel_window.minimize()
        if (pgw.getActiveWindowTitle() == "Microsoft Excel"):
            time.sleep(1)
            KMS.rightArrow()
            time.sleep(1)
            KMS.pressEnter()
        active_window = pgw.getActiveWindow()
        if active_window is not None:
            active_window.minimize()
        if stop_threads:
            break


def main(Name, Start_date):
    projects = []
    Thématiques_inconnues = []
    Thématiques_non_applicables = []
    NT_value = ''
    NA_value = ''
    print("hi")
    try:
        path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
        print("path ---->", path)
        Campagnec_Book = EI.openExcel(path)
        Campagnec_Book.activate()
        Check_list_sTr_sheet = Campagnec_Book.sheets['Check-list Start']
        sheet_value = Check_list_sTr_sheet.used_range.value
        config_name = EI.searchDataInColCache(sheet_value, 1, 'Config')
        print("config_name--------->", config_name)

        Fun_name14 = EI.searchDataInColCache(sheet_value, 1, 'Writter of the campaign')
        row, col = Fun_name14['cellPositions'][0]
        print(row, col)
        EI.setDataInCell(Check_list_sTr_sheet, (row, col + 4), Name)
        EI.setDataInCell(Check_list_sTr_sheet, (row + 1, col + 4), Start_date)
        data = EI.searchDataInColCache(sheet_value, 4, 'NI')
        data1 = EI.searchDataInColCache(sheet_value, 4, 'Global result')
        print("data1--------->", data1['cellPositions'])
        print("data--------->", data['cellPositions'])
        filtered_data = [item for i, item in enumerate(data['cellPositions']) if i in (1, 2, 3, 6, 7)]
        print(filtered_data)
        for position in filtered_data:
            print('position1-------->', position)
            row, col = position
            EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 1), 'Yes')

        for n, tuple_data in enumerate(config_name['cellPositions'][1:]):
            print(tuple_data)
            row, col = tuple_data
            # row, col = config_name['cellPositions'][1]
            print(row, col)
            if EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 4)) != 'No' and EI.getDataFromCell(
                    Check_list_sTr_sheet, (row, col + 4)) != 'NI':
                Type_of_Validation = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 4))
                Configurations = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 1))
                Priority = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 5))
                Name_of_Project = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                print("DFGSFDAFGAERF------->", Configurations, Type_of_Validation, Priority, Name_of_Project)

                EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 5), n + 1)
                if str(n+1) == str(1):
                    EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 4), 'PC1')
                    Config_sheet = 'Campagne Config 1'
                    print("Config_sheet1-------->", Config_sheet)
                    project1= EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                    projects.append(project1.split('(')[0].strip())
                    Thématiques_inconnue, Thématiques_non_applicable, NT_value, NA_value = getthematics(Config_sheet)
                    Thématiques_inconnues.append(Thématiques_inconnue)
                    Thématiques_non_applicables.append(Thématiques_non_applicable)

                else:
                    n = n-1
                    EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 4), 'Compl. PC'+str(n + 1))
                    Config_sheet = 'Campagne Config '+str((n+1)+1)
                    print("Config_sheet2-------->",Config_sheet)
                    project2 = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                    Thématiques_inconnue, Thématiques_non_applicable, NT_value, NA_value = getthematics(Config_sheet)
                    Thématiques_inconnues.append(Thématiques_inconnue)
                    Thématiques_non_applicables.append(Thématiques_non_applicable)
                    projects.append(project2.split('(')[0].strip())
        print("projects.append(project1)---------->", projects)
        print("1stsheet--------->", Thématiques_inconnues, Thématiques_non_applicables)

    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return projects, Thématiques_inconnues, Thématiques_non_applicables, NT_value, NA_value


def getthematics(Config_sheet):
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
    print("path ---->", path)
    Campagnec_Book = EI.openExcel(path)
    Campagnec_Book.activate()
    Campagne_sheet = Campagnec_Book.sheets[Config_sheet]
    sheet_value = Campagne_sheet.used_range.value
    Fun_name4 = EI.searchDataInColCache(sheet_value, 1, 'Function(s) or type of test')
    row, col = Fun_name4['cellPositions'][0]
    time.sleep(2)
    print(row, col)
    Fun_name4_1 = EI.searchDataInColCache(sheet_value, 5, 'Thématiques inconnues')
    row, col1 = Fun_name4_1['cellPositions'][0]
    time.sleep(2)
    print("row, col0000------------------>", row, col1)

    Thématiques_inconnues = EI.getDataFromCell(Campagne_sheet, (row, col1 + 1))
    print("Thématiques_inconnues---->", Thématiques_inconnues)

    Fun_name4_2 = EI.searchDataInColCache(sheet_value, 5, 'Thématiques non applicables')
    row, col2 = Fun_name4_2['cellPositions'][0]
    time.sleep(2)
    print("row, col0000------------------>", row, col2)
    Thématiques_non_applicables = EI.getDataFromCell(Campagne_sheet, (row, col2 + 1))
    print("Thématiques_non_applicables---->", Thématiques_non_applicables)

    NT_values = EI.getDataFromCell(Campagne_sheet, (row, col2 + 3))
    print("NT_values------>",NT_values)
    NA_values = EI.getDataFromCell(Campagne_sheet, (row+1, col2 + 3))
    print("NA_values------>", NA_values)

    NT_value = int(re.search(r':\s*(\d+)', NT_values).group(1)) if re.search(r':\s*(\d+)', NT_values) else None
    NA_value = int(re.search(r':\s*(\d+)', NA_values).group(1)) if re.search(r':\s*(\d+)', NA_values) else None

    print("NT_value------->", NT_value, NA_value)

    return Thématiques_inconnues, Thématiques_non_applicables, NT_value, NA_value


def getrequired_software():
    required_software = ''
    software = 'NEA R1.2 SP3 V6 BR2'
    # Define a regular expression pattern to match "NEA R1.2"
    pattern = r'(NEA R1\.2)'

    # Use re.search to find the pattern in the input string
    match = re.search(pattern, software)

    # Check if the pattern is found
    if match:
        # Extract the matched group
        required_software = match.group(1)
        print(required_software)
    else:
        print("Pattern not found.")
    return required_software


def getrequired_project_ID(projects):
    project_ID = ''
    input_string = projects[0]
    # input_string = "Gene_R1_2_D85_BEV_Manu"
    # Define a regular expression pattern to match a single letter followed by two digits
    pattern = r'([A-Za-z]\d{2}|[A-Za-z]{2}\d{1}|[A-Za-z]{1}\d{1}[A-Za-z]{1}|d{1}[A-Za-z]{1}|d{2}[A-Za-z]{1})'
    # Use re.search to find the pattern in the input string
    match = re.search(pattern, input_string)
    # Check if the pattern is found
    if match:
        # Extract the matched group
        project_ID = match.group(0)
        print(project_ID)
    else:
        print("Pattern not found.")
    return project_ID


def getthematics_in_list(Thématiques_non_applicables):
    # Thématiques_non_applicables = "| ABY_02 | AEM_01 | AFC_04 | AGG_01 | AMD_02 | AMG_01 | APT_04 | AWB_01 | AXG_01 | AXW_02 | AZX_02 | BPR_01 | CAE_04 | CAE_05 | CMP_02 | CRA_01 | CWF_03 | D12_01 | D41_01 | D41_02 | D59_02 | D7A_02 | D7K_01 | D7U_02 | D7U_04 | D7W_00 | D7W_03 | DAO_01 | DE7_03 | DI6_04 | DO8_00 | DRG_40 | DUE_07 | DUE_08 | DVQ_64 | DVQ_67 | DXD_00 | DXD_03 | DXD_05 | ENK_00 | EPD_00 | ICS_07 | IPD_00 | IWA_01 | IWH_00 | IWW_00 | LVM_02 | LXA_01"
    Thématiques_non_applicables = Thématiques_non_applicables.strip('|')
    # Split the string into a list and remove leading/trailing whitespaces from each element
    thematiques_list = [theme.strip() for theme in Thématiques_non_applicables.split('|')]
    # Remove any empty strings from the list
    thematiques_lists = list(filter(None, thematiques_list))
    # Print the list
    print(thematiques_lists)
    return thematiques_lists


def not_thematic(projects, Thématiques_non_applicables):
# def not_thematic():
    project = projects[0]
    Thématiques_non_applicable = Thématiques_non_applicables[0]
    # project = 'R1_2_D85_Gene_Manu'
    print("ko")
    result_tuples = []
    thematiques_lists = getthematics_in_list(Thématiques_non_applicable)
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[21]
    # print("path ---->", path)
    Silhouette_Book = EI.openExcel(path)
    Silhouette_Book.activate()
    Silhouettes_sheet = Silhouette_Book.sheets['Silhouettes']
    sheet_value = Silhouettes_sheet.used_range.value
    try:
        for result_string in thematiques_lists:
            print("thematiques_lists result_string------>", result_string)
            cleaned_result_string = result_string.strip()
            Fun_name9 = EI.searchDataInColCache(sheet_value, 1, cleaned_result_string)
            # print("Fun_name9----->",Fun_name9)
            # Check if 'cellPositions' is not empty before accessing its elements
            if Fun_name9['cellPositions']:
                row, col = Fun_name9['cellPositions'][0]
                time.sleep(2)
                # print(row, col)
                Project_cell = EI.searchDataInExcelCache(sheet_value, project)
                # print("project cells----------->", Project_cell)
                row18, col18 = Project_cell['cellPositions'][0]
                project1 = EI.getDataFromCell(Silhouettes_sheet, (row18, col18))
                # print("project1------------->", project1)
                if project == project1:
                    # print("hiiiiiiiii")
                    Silhouettes_sheet_Valeur = EI.getDataFromCell(Silhouettes_sheet, (row, col18))
                    print("ValeurValeurValeurValeur--->", Silhouettes_sheet_Valeur)
                    if Silhouettes_sheet_Valeur == '--' or Silhouettes_sheet_Valeur == 'X' or Silhouettes_sheet_Valeur == 'opt':
                        print("v--------------->", result_string)
                        result_tuples.append((cleaned_result_string, Silhouettes_sheet_Valeur))
        print("result_tuples---------------->", result_tuples)
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return result_tuples


def find_sheet_in_plm_book(PLM_EE_Book, project_ID, required_software):
    for sheet in PLM_EE_Book.sheets:
        sheet_name_parts = sheet.name.split('_')
        if len(sheet_name_parts) == 3:
            sheet_project_ID = sheet_name_parts[1]
            sheet_required_software = sheet_name_parts[2]
            if sheet_project_ID == project_ID and sheet_required_software == required_software:
                return sheet
    return None


def checkinplm(result_tuples, projects, NT_value, NA_value):
    # result_tuples =[('ABY_02', '--'), ('AGG_01', '--'), ('AMD_02', '--'), ('AMG_01', '--'),
    #                                   ('AWB_01', '--'), ('AXG_01', '--'), ('AXW_02', '--'), ('AZX_02', '--'),
    #                                   ('CMP_02', '--'), ('CWF_03', '--'), ('D41_01', '--'), ('D41_02', '--'),
    #                                   ('D59_02', '--'), ('D7A_02', '--'), ('D7K_01', '--'), ('D7U_02', '--'),
    #                                   ('D7U_04', '--'), ('DE7_03', '--'), ('DI6_04', '--'), ('DO8_00', '--'),
    #                                   ('DRG_40', '--'), ('DUE_07', '--'), ('DUE_08', '--'), ('DVQ_64', '--'),
    #                                   ('DVQ_67', '--'), ('EPD_00', '--'), ('IPD_00', '--'), ('IWA_01', '--'),
    #                                   ('IWH_00', '--'), ('IWW_00', '--'), ('LVM_02', '--'), ('LXA_01', '--')]

    NA_content = []
    project_ID = getrequired_project_ID(projects)
    required_software = getrequired_software()
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[22]
    print("path ---->", path)
    global stop_threads
    stop_threads = False
    t1 = threading.Thread(target=excel_popup, args=(EI.findInputFiles()[22],))
    t1.start()
    PLM_EE_Book = EI.openExcel(path)
    stop_threads = True
    PLM_EE_Book.activate()
    matching_sheet = find_sheet_in_plm_book(PLM_EE_Book, project_ID, required_software)
    if matching_sheet:
        print(f"Found sheet: {matching_sheet.name}")
    else:
        print("Sheet not found.")
    #
    # sheet = 'ENV_' + project_ID + '_' + required_software
    # print("sheet------>", sheet)

    try:
        func_sheet = PLM_EE_Book.sheets[matching_sheet]
        # Convert the list to an Excel range object
        sheet_range = func_sheet.range('A1').expand()
        # Extract values from the Excel range
        sheet_value = sheet_range.value
        # Find the last non-empty row
        maxrow = sheet_range.end('up').row
        print("maxrow--------->", maxrow)
        sheet_value = func_sheet.used_range.value
        Fun_name9 = EI.searchDataInColCache(sheet_value, 7, 'Name')
        print("Fun_name9----->", Fun_name9)
        # Check if 'cellPositions' is not empty before accessing its elements
        if Fun_name9['cellPositions']:
            row, col = Fun_name9['cellPositions'][0]
            time.sleep(2)
            print(row, col)
            for i in result_tuples:
                if i[1] == '--' or i[1] == 'X':
                    print("i--------->", i[0])
                    result_tuple = i[0]
                    print("result_tuple.split('_')[0]---->",result_tuple.split('_')[0])
                    try:
                        # PLM_sheet_Valeur = EI.searchDataInExcelCache(sheet_value, (1, maxrow), result_tuple.split('_')[0])
                        PLM_sheet_Valeur = EI.searchDataInColCache(sheet_value, col, result_tuple.split('_')[0])
                        print("PLM_sheet_Valeur--------->", PLM_sheet_Valeur)
                        if PLM_sheet_Valeur['count'] == 0:
                            print("PLM_sheet_Valeur['count'] == 0")
                            b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA'
                            NA_content.append(b)
                        elif PLM_sheet_Valeur['count'] == 1:
                            print("PLM_sheet_Valeur['count'] == 1")
                            row, col = PLM_sheet_Valeur['cellPositions'][0]
                            thematic_value = EI.getDataFromCell(func_sheet, (row, col+2))
                            b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is fixed on the value ' + thematic_value
                            NA_content.append(b)
                        elif PLM_sheet_Valeur['count'] >= 2:
                            print("PLM_sheet_Valeur['count'] > 2")
                            result_list = []
                            for cellposition in PLM_sheet_Valeur['cellPositions']:
                                row, col = cellposition
                                print("row,col", row, col)
                                thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is fixed on the value ' + thematic_value
                                NA_content.append(b)
                            #     result_list.append(thematic_value)
                            # result = "|".join(result_list)
                            # b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is fixed on the value ' + result
                            # NA_content.append(b)

                    except Exception as ex:
                        exc_type, exc_obj, exc_tb = sys.exc_info()
                        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
                        print("Thematics are not present in the PLM_Sheet.")
                if i[1] == 'opt':
                    print("i--------->", i[0])
                    result_tuple = i[0]
                    print("result_tuple.split('_')[0]---->", result_tuple.split('_')[0])
                    try:
                        PLM_sheet_Valeur = EI.searchDataInColCache(sheet_value, col, result_tuple.split('_')[0])
                        print("PLM_sheet_Valeur--------->", PLM_sheet_Valeur)
                        if PLM_sheet_Valeur['count'] > 0:
                            print("PLM_sheet_Valeur['count'] > 0")
                            result_list = []
                            for cellposition in PLM_sheet_Valeur['cellPositions']:
                                row, col = cellposition
                                print("row,col", row, col)
                                thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                result_list.append(thematic_value)
                            result = "|".join(result_list)
                            b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is optional on the value ' + result
                            NA_content.append(b)
                    except Exception as ex:
                        exc_type, exc_obj, exc_tb = sys.exc_info()
                        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
                        print("Thematics are not present in the PLM_Sheet.")


    except:
        print("sheet are Thematics are not present in the PLM_Sheet.")

    project_text_bold = "\033[1m" + projects[0] + "\033[0m"
    print(project_text_bold)
    # doc_content = [projects[0] + ':', '\n ', '\n ', 'Total number of NT files: 178','Total number of NA files: 285', '\n ', ' \n', 'NA because of following thematiques: ', '\n ', '\n']
    doc_content = ['Total number of NT files: '+str(NT_value), 'Total number of NA files: '+str(NA_value), 'NA because of following thematiques: ']
    NA_NT_content = doc_content + NA_content
    # NA_NT_content = NA_content
    print("NA_NT_content-------->", NA_NT_content)
    return NA_NT_content


# def append_to_word_document(content_list, output_path):
#     # Create a new Word document
#     doc = Document()
#     # Append each element in the content list to the document
#     for element in content_list:
#         # Split the element by ':'
#         parts = element.split(':')
#         if len(parts) == 2:
#             # Add both parts on the same line
#             doc.add_paragraph(f"{parts[0].strip()}: {parts[1].strip()}")
#         else:
#             # If there is no ':', add the element as a paragraph
#             doc.add_paragraph(element.strip())
#     # Save the document
# #     doc.save(output_path)
# def append_to_word_document(content, output_path):
#     doc = Document()
#
#     for line in content:
#         sanitized_line = ''.join(c for c in line if c.isprintable())
#
#         # Check if the line is not empty after sanitization
#         if sanitized_line.strip():
#             doc.add_paragraph(sanitized_line.strip())
#
#     doc.save(output_path)


def append_to_word_document(content, output_path):
    doc = Document()

    # Assume the first element is the heading
    heading = content[0]
    # Add the heading to the document with bold formatting
    doc.add_paragraph(heading.strip(), style='Heading1')

    # Iterate through the rest of the content
    for line in content[1:]:
        # Treat '\n' as a new line in the Word document
        if line.strip() == '\n':
            doc.add_paragraph()
        else:
            sanitized_line = ''.join(c for c in line if c.isprintable())

            # Check if the line is not empty after sanitization
            if sanitized_line.strip():
                doc.add_paragraph(sanitized_line.strip())

    doc.save(output_path)


def updateFeps(projects, Fepss):
# def updateFeps():
#     project = 'R1_2_D85_Gene_Manu'
#     Fepss = ['FEPS_117176','FEPS_117177','FEPS_117218','FEPS_119754','FELP_120108','FEPS_119753','FELP_120104','FEPS_119956','FEPS_119954']
    project = projects[0]
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
    print("path ---->", path)
    Campagnec_Book = EI.openExcel(path)
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[1]
    print("path ---->", path)
    test_Book = EI.openExcel(path)
    contents = []
    for Feps in Fepss:
        numeric_part = ''.join(filter(str.isdigit, Feps))
        print("numeric_part--------->",numeric_part)
        NA_NT_values = []
        try:
            test_Book.activate()
            Feps_sheet = test_Book.sheets['FEPS History']
            sheet_value = Feps_sheet.used_range.value
            Fun_name4 = EI.searchDataInColCache(sheet_value, 3, numeric_part.strip())
            row, col = Fun_name4['cellPositions'][0]
            time.sleep(2)
            print(row, col)
            Impacted_sheets = EI.getDataFromCell(Feps_sheet, (row, col-1))
            print(f'Impacted_sheets for the FEPS_{numeric_part}---------->{Impacted_sheets}')
            Campagnec_Book.activate()
            Campagne_sheet = Campagnec_Book.sheets['Synthèse des campagnes']
            sheet_value = Campagne_sheet.used_range.value
            # Assuming Impacted_sheets is a list of sheet names
            print("Impacted_sheets.split()--------->", Impacted_sheets.split('\n'))
            impacetdSheets = Impacted_sheets.split('\n')
            for sheet_name in impacetdSheets:
                if sheet_name:
                    print(f'Do something with sheet: {sheet_name}')
                    Fun_name4 = EI.searchDataInColCache(sheet_value, 1, sheet_name)
                    print("Fun_name4---------->", Fun_name4)
                    if Fun_name4['cellPositions']:
                        for cellPositions in Fun_name4['cellPositions']:
                            row, col5 = cellPositions
                            print("row, col5----------->", row, col5)
                            EI.setDataFromCell(Campagne_sheet, (row, col5+12), Feps)
                            project1 = EI.getDataFromCell(Campagne_sheet, (15, 8))
                            if project == project1:
                                print("hiiiiiiiii")
                                NA_NT_value = EI.getDataFromCell(Campagne_sheet, (row, 8))
                                print("NA_NT_values------->",NA_NT_value)
                                NA_NT_values.append(NA_NT_value)
            print("NA_NT_values-->", NA_NT_values)
            count_NT = NA_NT_values.count('NT')
            count_NA = NA_NT_values.count('NA')
            content = 'For '+Feps+' will be '+ str(count_NT) +' NT and '+str(count_NA)+' NA'
            contents.append(content)
        except Exception as ex:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
            print("FEps are not present in the testplan FEPS History sheet.")
    project_text_bold =  "\033[1m" + project + "\033[0m"
    print(project_text_bold)
    # doc_content = [project_text_bold+':', ' ', ' ']
    # doc_content = [project+':', '\n ', '\n']
    doc_content = [project+':']
    NA_NT_content = doc_content + contents
    # NA_NT_content = contents
    print("contentscontentscontents---------->", NA_NT_content)
    return NA_NT_content


if __name__ == '__main__':
    ICF.loadConfig()
    print("Tool Started")
    start = time.time()
    # FEPS_117176,FEPS_117177,FEPS_117218,FEPS_119754,FELP_120108,FEPS_119753,FELP_120104,FEPS_119956,FEPS_119954
    Name = "RAMA KRISHNA REDDY KOPPULA (RRE - EXP)"
    Software = "SP3_V6.00 Official"
    Start_date = "12/28/2023"
    Fepss = input("Enter the GAELE_reference by giving , : ")
    Fepss = Fepss.split(',')
    projects, Thématiques_inconnues, Thématiques_non_applicables, NT_value, NA_value = main(Name, Start_date)
    no_of_NA_NT_contents = updateFeps(projects, Fepss)
    result_tuples = not_thematic(projects, Thématiques_non_applicables)
    # result_tuples = not_thematic()
    NA_content = checkinplm(result_tuples, projects, NT_value, NA_value)
    # Example usage
    output_path = r'C:\Users\vgajula\Downloads\output_document.docx'
    ALl_content = no_of_NA_NT_contents + NA_content

    print("ALl_content-------->",ALl_content)
    append_to_word_document(ALl_content, output_path)
    print(f"Word document saved to: {output_path}")
    end1 = time.time()
    print("\nexecution time " + str(end1 - start))
    print("Tool Completed")





