# # # # import docx
# # # # import re
# # # # import difflib
# # # # import ctypes
# # # # import concurrent.futures
# # # # from concurrent.futures import ThreadPoolExecutor
# # # # import ExcelInterface as EI
# # # # import json
# # # # import os
# # # # from os import listdir
# # # # from os.path import isfile, join
# # # # import sys
# # # # import shutil
# # # # import QIAParamCreateNewFrame as QIAP
# # # # import logging
# # # # import WordDocInterface as WD
# # # # import os
# # # # import InputConfigParser as ICF
# # # #
# # # #
# # # # def getContent(Doc,ReqName,ReqVer):
# # # #     TableList= WD.getTables(Doc)
# # # #     #RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
# # # #     RqTable= WD.threading_findTable(TableList, ReqName)
# # # #     logging.info("")
# # # #     if RqTable==-1:
# # # #         if ((ReqName.find('.')!=-1)|(ReqName.find('_')!=-1)):
# # # #             ReqName=ReqName.replace('.', '-')
# # # #         RqTable= WD.threading_findTable(TableList, ReqName+"("+ReqVer+")")
# # # #     else:
# # # #         RqTable= WD.threading_findTable(TableList, ReqName+"("+ReqVer+")")
# # # #     if RqTable!=-1:
# # # #         chkOldFormat= WD.checkFormat(RqTable, ReqName+"("+ReqVer+")")
# # # #         if chkOldFormat==0:
# # # #             Content= WD.getOldContents(RqTable, ReqName+"("+ReqVer+")")
# # # #         else:
# # # #             Content= WD.getNewContents(RqTable, ReqName+"("+ReqVer+")")
# # # #     else:
# # # #         RqTable= WD.threading_findTable(TableList, ReqName+" "+ReqVer)
# # # #         if RqTable!=-1:
# # # #             chkOldFormat= WD.checkFormat(RqTable, ReqName+" "+ReqVer)
# # # #             if chkOldFormat==0:
# # # #                 Content= WD.getOldContents(RqTable, ReqName+" "+ReqVer)
# # # #             else:
# # # #                 Content= WD.getNewContents(RqTable, ReqName+" "+ReqVer)
# # # #         else:
# # # #             RqTable= WD.threading_findTable(TableList, ReqName+"  "+ReqVer)
# # # #             if RqTable!=-1:
# # # #                 chkOldFormat= WD.checkFormat(RqTable, ReqName+"  "+ReqVer)
# # # #                 if chkOldFormat==0:
# # # #                     Content= WD.getOldContents(RqTable, ReqName+"  "+ReqVer)
# # # #                 else:
# # # #                     Content= WD.getNewContents(RqTable, ReqName+"  "+ReqVer)
# # # #             else:
# # # #                 Content = -1
# # # #     logging.info("content123--->",Content)
# # # #     return Content
# # # #
# # # # def getReqVer(req):
# # # #     if req.find('(') != -1:
# # # #         new_reqName = req.split("(")[0].split()[0] if len(req.split("(")) > 0 else ""
# # # #         new_reqVer = req.split("(")[1].split(")")[0] if len(req.split("(")) > 1 else ""
# # # #     else:
# # # #         new_reqName = req.split()[0] if len(req.split()) > 0 else ""
# # # #         new_reqVer = req.split()[1] if len(req.split()) > 1 else ""
# # # #     return new_reqName.strip(), new_reqVer.strip()
# # # #
# # # #
# # # # def confirmationPopup(func):
# # # #     global confirmationPop
# # # #     confirmationPop = func
# # # #     return confirmationPop
# # # #
# # # # # didInfo = {'req': requirement, 'didVal': DID_value, 'cerebro': "cerebro.inetpsa.com/data-identifier-list"}
# # # # #                             DID_type = confirmationPop(didInfo)
# # # #
# # # #
# # # # if __name__=="__main__":
# # # #     ICF.loadConfig()
# # # #     Doc = "C:\\Users\\vgajula\\Documents\\07-11-2023\\task surya\\Input_Files\\Output_Folder\\[V18.0][02017_19_02196]SSD_HMIF_LONGITUDINAL_MOBILITY_MOBY_HMI_23Q3.docx"
# # # #     req = "REQ-0854846 A"
# # # #     ReqName, ReqVer = getReqVer(req)
# # # #     doc = getContent(Doc, ReqName, ReqVer)
# # # #     print("content-------->",doc)
# # # #     ID_type = confirmationPopup(doc)
# # # #
# # #
# # #
# # #
# # # import os
# # # from docx import Document  # Assuming you have the python-docx library installed
# # # import WordDocInterface as WDI
# # # from selenium import webdriver
# # # from selenium.common import exceptions
# # # from selenium.webdriver.common.keys import Keys
# # # import time
# # # from selenium.webdriver.chrome.service import Service
# # # from selenium.webdriver.common.by import By
# # # from selenium.webdriver.support.ui import WebDriverWait
# # # from selenium.webdriver.support import expected_conditions as EC
# # # import shutil
# # # import os
# # # import InputConfigParser as ICF
# # # import logging
# # #
# # # def search_keyword_in_table(table, keyword):
# # #     text = ''
# # #     if table != -1:
# # #         i=0
# # #         for col in table.columns:
# # #             j=0
# # #             for cell in col.cells:
# # #                 if len(cell.text) !=0:
# # #                     if keyword in cell.text:
# # #                         text=table.columns[i+1].cells[j].text
# # #                 j=j+1
# # #             i=i+1
# # #     else:
# # #         text=-1
# # #     return text
# # #
# # # def process_word_document(input_path, output_path, keyword):
# # #     # Open the document
# # #     doc = Document(input_path)
# # #
# # #     # Perform your operations on the document here
# # #     keyword_found = search_keyword_in_table(doc, keyword)
# # #
# # #     # Save the modified document to the output folder
# # #     doc.save(output_path)
# # #
# # #     return keyword_found
# # #
# # # def find_and_save_keyword(file_path, keyword, output_folder):
# # #     try:
# # #         # Open the Word document
# # #         doc = Document(file_path)
# # #
# # #         # Check if the keyword is present in the document
# # #         if any(keyword in paragraph.text for paragraph in doc.paragraphs):
# # #             # Create the output folder if it doesn't exist
# # #             if not os.path.exists(output_folder):
# # #                 os.makedirs(output_folder)
# # #
# # #             # Extract the filename from the original path
# # #             file_name = os.path.basename(file_path)
# # #
# # #             # Create the output file path
# # #             output_path = os.path.join(output_folder, file_name)
# # #
# # #             # Save the document to the output folder
# # #             doc.save(output_path)
# # #
# # #             print(f"Keyword found in '{file_name}'. Document saved to '{output_folder}'.")
# # #     except Exception as e:
# # #         print(f"Error processing '{file_path}': {e}")
# # #
# # # def process_folder(input_folder, keyword, output_folder):
# # #     # Get a list of all files in the input folder
# # #     files = [f for f in os.listdir(input_folder) if f.endswith(".docx")]
# # #
# # #     # Process each file in the input folder
# # #     for file in files:
# # #         file_path = os.path.join(input_folder, file)
# # #         find_and_save_keyword(file_path, keyword, output_folder)
# # #
# # #
# # # def getContent(Doc,keyword):
# # #     # def getContent(Doc,ReqName,ReqVer):
# # #     try:
# # #         TableList = WDI.getTables(Doc)
# # #         #RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
# # #         RqTable = WDI.threading_findTable(TableList, keyword)
# # #         logging.info("")
# # #         if RqTable == -1:
# # #             if keyword!=-1:
# # #                 RqTable = WDI.threading_findTable(TableList, keyword)
# # #         else:
# # #             RqTable = WDI.threading_findTable(TableList, keyword)
# # #         if RqTable != -1:
# # #             chkOldFormat = WDI.checkFormat(RqTable, keyword)
# # #             if chkOldFormat == 0:
# # #                 Content = WDI.getOldContents(RqTable, keyword)
# # #             else:
# # #                 Content = WDI.getNewContents(RqTable, keyword)
# # #     except:
# # #         pass
# # #     # else:
# # #     #     RqTable = WDI.threading_findTable(TableList, keyword)
# # #     #     if RqTable != -1:
# # #     #         chkOldFormat = WDI.checkFormat(RqTable, ReqName+" "+ReqVer)
# # #     #         if chkOldFormat == 0:
# # #     #             Content = WDI.getOldContents(RqTable, ReqName+" "+ReqVer)
# # #     #         else:
# # #     #             Content = WDI.getNewContents(RqTable, ReqName+" "+ReqVer)
# # #     #     else:
# # #     #         RqTable = WDI.threading_findTable(TableList, ReqName+"  "+ReqVer)
# # #     #         if RqTable != -1:
# # #     #             chkOldFormat = WDI.checkFormat(RqTable, ReqName+"  "+ReqVer)
# # #     #             if chkOldFormat == 0:
# # #     #                 Content = WDI.getOldContents(RqTable, ReqName+"  "+ReqVer)
# # #     #             else:
# # #     #                 Content = WDI.getNewContents(RqTable, ReqName+"  "+ReqVer)
# # #     #         else:
# # #     #             Content = -1
# # #     logging.info("content123--->", Content)
# # #     return Content
# # #
# # #
# # # if __name__ == "__main__":
# # #     input_folder = "C:\\Users\\vgajula\\Documents\\07-11-2023\\task surya\\Input_Files"
# # #     output_folder = "C:\\Users\\vgajula\\Documents\\07-11-2023\\task surya\\Input_Files\\Output_Folder"
# # #     keyword = "BRAKE_REQUEST"
# # #
# # #     # Check if the output folder exists, and create it if not
# # #     if not os.path.exists(output_folder):
# # #         os.makedirs(output_folder)
# # #
# # #     process_folder(input_folder, keyword, output_folder)
# # #
# # #     #   WDI.getContent(Doc,ReqName,ReqVer)
# # #     # Loop through each file in the input folder
# # #     for filename in os.listdir(input_folder):
# # #         if filename.endswith(".docx") or filename.endswith(".doc"):  # Check if it's a Word document
# # #             input_path = os.path.join(input_folder, filename)
# # #             output_path = os.path.join(output_folder, filename)
# # #
# # #             # Process the Word document and check if the keyword is in a table cell starting with "REQ-"
# # #             keyword_found = process_word_document(input_path, output_path, keyword)
# # #
# # #             if keyword_found:
# # #                 print(f"Keyword '{keyword_found}' found in {filename}")
# # #     print("Processing complete.")
# # #
# # #
# # #
# # #
# # # # def getNewContents(table, keyword):
# # # #     logging.info("In getNewContents function finding cnt")
# # # #     text = -1
# # # #     data = {'req': keyword}
# # # #     if table != -1:
# # # #         i = 0
# # # #         for col in table.columns:
# # # #             j = 0
# # # #             for cell in col.cells:
# # # #
# # # #                 if len(cell.text) != 0:
# # # #                     # logging.info("t",cell.text)
# # # #                     if "n°requirement" in cell.text.lower():
# # # #                         WDI.clearStrikethrough(table.columns[i].cells[j + 1])
# # # #                         text = (table.columns[i].cells[j + 1].text.encode('utf-8').strip())
# # # #                         text = text.decode('utf-8').strip()
# # # #                         oldName = re.findall(r'GEN-.*\)', text)
# # # #                         newName = re.findall(r'REQ-.*', text)
# # # #                         data["oldName"] = oldName
# # # #                         data['req'] = newName
# # # #                     if "content of the requirement" in cell.text.lower():
# # # #                         WDI.clearStrikethrough(table.columns[i].cells[j + 1])
# # # #                         text = (table.columns[i].cells[j + 1].text.encode('utf-8').strip())
# # # #                         text = text.decode('utf-8').strip()
# # # #                         data["content"] = text
# # # #                         # logging.info("new text2",text)
# # # #                     elif "effectivity expression" in cell.text.lower():
# # # #                         WDI.clearStrikethrough(table.columns[i].cells[j + 1])
# # # #                         text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
# # # #                         text = text.decode('utf-8').strip()
# # # #                         data["effectivity"] = text
# # # #                     elif "lcdv" in cell.text.lower():
# # # #                         WDI.clearStrikethrough(table.columns[i].cells[j + 1])
# # # #                         text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
# # # #                         text = text.decode('utf-8').strip()
# # # #                         data["lcdv"] = text
# # # #                     elif "diversity" in cell.text.lower():
# # # #                         WDI.clearStrikethrough(table.columns[i].cells[j + 1])
# # # #                         text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
# # # #                         text = text.decode('utf-8').strip()
# # # #                         data["diversity"] = text
# # # #                     elif "target configuration" in cell.text.lower():
# # # #                         WDI.clearStrikethrough(table.columns[i].cells[j + 1])
# # # #                         text = (table.columns[i + 1].cells[j].text.encode('utf-8').strip())
# # # #                         text = text.decode('utf-8').strip()
# # # #                         data["target configuration"] = text
# # # #                 j = j + 1
# # # #             i = i + 1
# # # #     else:
# # # #         data = -1
# # # #     # logging.info("In getNewContents function returning text",text)
# # # #     return data
# #
# #
# # from docx import Document
# #
# # def find_tables_by_keyword_and_header(doc, keyword, header):
# #     matching_tables = []
# #
# #     for table in doc.tables:
# #         # Check if the first cell in the first row contains the specified header
# #         if table.cell(0, 0).text.strip().lower() == header.lower():
# #             for row in table.rows:
# #                 for cell in row.cells:
# #                     if keyword in cell.text:
# #                         matching_tables.append(table)
# #                         break
# #                 if matching_tables:
# #                     break
# #
# #     return matching_tables
# #
# #
# # def clearStrikethrough(cell):
# #     for paragraph in cell.paragraphs:
# #         for run in paragraph.runs:
# #             # Check whether the run text is strikethrough
# #             if run.font.strike:
# #                 # Remove the strikethrough text
# #                 run.clear()
# #
# # def confirmationPopup(func):
# #     global confirmationPop
# #     confirmationPop = func
# #     return confirmationPop
# #
# # # didInfo = {'req': requirement, 'didVal': DID_value, 'cerebro': "cerebro.inetpsa.com/data-identifier-list"}
# # # DID_type = confirmationPop(didInfo)
# #
# # def print_table_content(table, keyword):
# #     keyword_found = False
# #
# #     # Check if the keyword is present in the table
# #     for row in table.rows:
# #         for cell in row.cells:
# #             if keyword in cell.text:
# #                 keyword_found = True
# #                 break
# #         if keyword_found:
# #             break
# #
# #     # If the keyword is found, print the table content
# #     if keyword_found:
# #         for row in table.rows:
# #             for cell in row.cells:
# #                 print(cell.text, end='\t')
# #                 # confirmationPopup(cell.text)
# #             print()
# #     else:
# #         print(f"No keyword '{keyword}' found in the table.")
# #
# # # Example usage:
# # # Assuming 'your_document.docx' is the Word document file, 'BRAKE_REQUEST' is the keyword, and 'n°requirement' is the header
# # doc_path = r'C:\Users\vgajula\Documents\07-11-2023\task surya\Input_Files\Keyword_file_folder\[V11.0][00998_17_02376]TFD_FT0_2_1_MVMS.docx'
# # keyword = 'BRAKE_REQUEST'
# # header = 'n°requirement'
# #
# # # Load the Word document
# # doc = Document(doc_path)
# # print("doc--------->",doc)
# # # Find the table containing the keyword and starting with the specified header
# # tables = find_tables_by_keyword_and_header(doc, keyword, header)
# # print("tables->",tables)
# #
# # for table in tables:
# #     # Print the content of the table if the keyword is found
# #     if table:
# #         print_table_content(table, keyword)
# #     else:
# #         print(f"No table found with the keyword '{keyword}' and starting with the header '{header}'.")
# #
# #
# #
# #
#
# import os
# import tkinter as tk
# from tkinter import scrolledtext
# from docx import Document
#
# def create_word_document(content, output_folder, output_filename):
#     # Create a new Word document
#     doc = Document()
#     doc.add_paragraph(content)
#
#     # Create the output folder if it doesn't exist
#     if not os.path.exists(output_folder):
#         os.makedirs(output_folder)
#
#     # Save the document to the output folder
#     output_path = os.path.join(output_folder, output_filename)
#     doc.save(output_path)
#
# def show_content_in_popup_and_save_to_doc(content, output_folder, output_filename):
#     # Create the main window
#     root = tk.Tk()
#     root.title("Table Content")
#
#     # Create a scrolled text widget for displaying the content
#     text_widget = scrolledtext.ScrolledText(root, width=80, height=20)
#     text_widget.insert(tk.END, content)
#     text_widget.configure(state='disabled')  # Make the text widget read-only
#
#     # Pack the text widget into the window
#     text_widget.pack()
#
#     # Button to save the content to a Word document
#     save_button = tk.Button(root, text="Save to Word", command=lambda: save_to_word(content, output_folder, output_filename))
#     save_button.pack()
#
#     # Start the Tkinter event loop
#     root.mainloop()
#
# def save_to_word(content, output_folder, output_filename):
#     create_word_document(content, output_folder, output_filename)
#
# if __name__ == "__main__":
#     # Example content
#     content = "[[['N°Requirement', 'Content of the Requirement'], ['REQ-0802334 B\nState : Distribuer\nModified : 9/25/2023 6:26:41 PM\nREQ-0802334\nOwner : se20228\nGeneric : Yes\nMaturity : Robust\nGrade : ASIL_QM\nKey Requirement : Yes\nFlexibility : \nAbstraction Level : \nPSA ISAF Compliance : Yes\nPSA SSTG : ', '(…) shall define the intern position of the shifter to the position D from N\n\nIF (\n   STABLE_LEVER_POS = N\n   AND\n   User_Transmission_Intent switches to D_Wanted\n)\n\nTHEN {\n   IF (\n      ETAT_MT = { MOTEUR_TOURNANT OR ARRETE}  // the state of the engine is running or stopped\n\n      AND [   // success conditions\n         CONTACT_FREIN = Pressed     // the brake is pressed\n         OR\n         VITESSE_VEHICULE_ROUES > PRM_V_SHIFT_LOCK     // The vehicle speed is superior than a threshold\n         OR   \n         POS_LEVIER_BV = N for less than PRM_TIME_POS_N   //  the gearshifter is on N for less than the associated temporization\n         OR\n         InhibitRequest = TRUE  // the auto park inhibition mode is activated\n      ]\n   )\n   THEN {\n      STABLE_LEVER_POS  switches from N to D\n   }\n\nELSE IF (   // refusal to pass from N to D for brake not pressed\n   CONTACT_FREIN <> Pressed\n   AND     // The vehicle speed is lower than a threshold\n   VITESSE_VEHICULE_ROUES lower or equal to  PRM_V_SHIFT_LOCK\n   AND // the gearshifter is on N for more than the associated temporization\n   POS_LEVIER_BV = N for more than PRM_TIME_POS_N\n)\nTHEN {\n   STABLE_LEVER_POS remains in N\n   AND\n   BRAKE_REQUEST = TRUE\n   AND \n   The PRM_TIME_CANCEL_BRAKE_REQUEST timer starts\n   AND \n   TransmissionRequestRefusalCause switches to PSF\n}\n\n'], ['Comments', 'The messages is for refusing the request of driver to pass from N to D when the brake pedal is not pressed'], ['Effectivity Expression', 'HP-0000873[L-∞] AND ( HP-0000873[AFE{AFE_08}] )'], ['LCDV Simplified EN', 'HP-0000873 [ DICO MULTIGAMME Q2_2023 - ∞ ] AND AFE TYPE_BV_LEVIER(AFE_08 ROTARY_SHIFTER_ENDLESS)'], ['Input Requirement', '']]]"
#
#     # Output folder and filename
#     output_folder = "C:\\Users\\vgajula\\Documents\\07-11-2023\\task surya\\Output_Files"
#     output_filename = "output_document.docx"
#
#     # Show the content in a popup and save to Word on button click
#     show_content_in_popup_and_save_to_doc(content, output_folder, output_filename)



import os
import tkinter as tk
from tkinter import scrolledtext
from docx import Document

def create_word_document(content, output_folder, output_filename):
    # Create a new Word document or open an existing one
    if os.path.exists(os.path.join(output_folder, output_filename)):
        doc = Document(os.path.join(output_folder, output_filename))
    else:
        doc = Document()

    # Add content to the document
    for section in content:
        for entry in section:
            doc.add_heading(entry[0], level=1)
            doc.add_paragraph(entry[1])

    # Save the document
    doc.save(os.path.join(output_folder, output_filename))

def show_content_in_popup_and_save_to_doc(content, output_folder, output_filename):
    # Convert the list to a string
    content_str = '\n'.join([' '.join(entry) for section in content for entry in section])

    # Create the main window
    root = tk.Tk()
    root.title("Table Content")

    # Create a scrolled text widget for displaying the content
    text_widget = scrolledtext.ScrolledText(root, width=80, height=20)
    text_widget.insert(tk.END, content_str)
    text_widget.configure(state='disabled')  # Make the text widget read-only

    # Pack the text widget into the window
    text_widget.pack()

    # Button to save the content to a Word document
    save_button = tk.Button(root, text="Save to Word", command=lambda: save_to_word(content, output_folder, output_filename))
    save_button.pack()

    # Start the Tkinter event loop
    root.mainloop()

def save_to_word(content, output_folder, output_filename):
    create_word_document(content, output_folder, output_filename)

if __name__ == "__main__":
    # Example content
    content = [[['N°Requirement', 'Content of the Requirement'],
                ['REQ-0854846 A\nState : Distribuer\nModified : 9/25/2023 6:26:41 PM\nREQ-0854846\nOwner : se20228\nGeneric : Yes\nMaturity : Robust\nGrade : \nKey Requirement : Yes\nFlexibility : \nAbstraction Level : \nPSA ISAF Compliance : Yes\nPSA SSTG : ',
                 '// This transition describes the required conditions to turn off a brake request.\nIF (\n   CONTACT_FREIN= Pressed\n   OR\n   STABLE_LEVER_POS value has changed\n   OR\n   The PRM_TIME_CANCEL_BRAKE_REQUEST timer ends\nTHEN {\n   BRAKE_REQUEST = FALSE\n}\n\n'],
                ['Comments', 'Equivalent to REQ-0578924 in AFE_06 without ACPK'],
                ['Effectivity Expression', 'HP-0000873[M-∞] AND ( HP-0000873[AFE{AFE_08}] )'],
                ['LCDV Simplified EN', 'HP-0000873 [ DICO MULTIGAMME Q3_2023 - ∞ ] AND AFE TYPE_BV_LEVIER(AFE_08 ROTARY_SHIFTER_ENDLESS)'],
                ['Input Requirement', '']]]

    # Output folder and filename
    output_folder = "output"
    output_filename = "output_document.docx"

    # Show the content in a popup and append to Word on button click
    show_content_in_popup_and_save_to_doc(content, output_folder, output_filename)
